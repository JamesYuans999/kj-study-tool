import streamlit as st
import requests
import json
import datetime
import pandas as pd
import pdfplumber
import time
import docx
import random
from supabase import create_client, ClientOptions
import plotly.express as px
from openai import OpenAI
import streamlit.components.v1 as components
import os
import edge_tts
import asyncio
import tempfile
import uuid
import re
import gc
import hashlib
import math
import unicodedata

# ==============================================================================
# 1. 全局配置与 CSS (紧急修复版：恢复原生交互)
# ==============================================================================
st.set_page_config(page_title="中级会计 AI 私教 Pro", page_icon="🥝", layout="wide", initial_sidebar_state="auto")

st.markdown("""
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.10.5/font/bootstrap-icons.css">
<style>
    /* =======================================
       1. 关键修复：恢复侧边栏原生行为
       ======================================= */
    /* 绝对不要设置 position: fixed !important，否则关不掉 */
    section[data-testid="stSidebar"] {
        background-color: #ffffff;
        border-right: 1px solid #f0f0f0;
        z-index: 99998 !important; /* 仅保证比 Header 高，但不锁定位置 */
    }

    /* =======================================
       2. 汉堡菜单按钮 (必须置顶)
       ======================================= */
    /* 确保按钮在最上层，否则点不到 */
    button[data-testid="stSidebarCollapsedControl"] {
        display: block !important;
        z-index: 999999 !important; /* 最高层级 */
        color: #00C090 !important; /* 绿色图标 */
        position: fixed; /* 按钮固定在左上角 */
        top: 10px;
        left: 10px;
        background: rgba(255,255,255,0.8); /*稍微加点背景防混淆*/
        border-radius: 50%;
        width: 2.5rem;
        height: 2.5rem;
    }

    /* 兼容旧版 ID */
    [data-testid="collapsedControl"] {
        display: block !important;
        z-index: 999999 !important;
        color: #00C090 !important;
        position: fixed;
        top: 10px;
        left: 10px;
    }

    /* =======================================
       3. 顶部 Header (防止遮挡)
       ======================================= */
    header[data-testid="stHeader"] {
        background: rgba(255, 255, 255, 0.95) !important;
        z-index: 99 !important; /* 比侧边栏低 */
        height: 3.75rem;
    }

    /* 隐藏彩虹条 */
    [data-testid="stDecoration"] { display: none !important; }

    /* =======================================
       4. 手机端内容避让 (关键)
       ======================================= */
    @media (max-width: 768px) {
        /* 强制给主内容区顶部加 padding，把内容“顶”下来 */
        .main .block-container {
            padding-top: 5rem !important; 
            max-width: 100vw !important;
        }
    }

    /* =======================================
       5. 视觉美化 (Bento Grid)
       ======================================= */
    .stApp { background-color: #F9F9F0; font-family: 'Segoe UI', sans-serif; }

    /* 卡片 */
    .css-card {
        background: #fff; border-radius: 12px; padding: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.03); border: 1px solid #eee; margin-bottom: 15px;
    }

    /* 统计数字 */
    .stat-value { font-size: 2rem; font-weight: 800; color: #333; }

    /* 按钮 */
    .stButton>button {
        background: #00C090; color: white; border: none; border-radius: 8px;
        height: 45px; font-weight: 600;
    }
    .stButton>button:hover { background: #00a87e; color: white; }

</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. 数据库连接与配置
# ==============================================================================
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
    SUPABASE_URL = st.secrets["supabase"]["url"]
    SUPABASE_KEY = st.secrets["supabase"]["key"]

    # 修复：仅在非 Streamlit Cloud 环境或显式要求时才启用代理
    # Streamlit Cloud 通常不需要代理即可访问 Supabase
    # 建议直接注释掉下面这几行，或者确保云端 Secrets 不包含 [env]
    # if "env" in st.secrets:
    #     os.environ["http_proxy"] = st.secrets["env"]["http_proxy"]
    #     os.environ["https_proxy"] = st.secrets["env"]["https_proxy"]
except:
    st.error("🔒 Secrets 配置丢失！请检查 .streamlit/secrets.toml 文件。")
    st.stop()


@st.cache_resource
def init_supabase():

    return create_client(
        SUPABASE_URL,
        SUPABASE_KEY,
        options=ClientOptions(
            postgrest_client_timeout=10,
            storage_client_timeout=10
        )
    )

supabase = init_supabase()

# 用户身份模拟 (生产环境需对接 st.login)
if 'user_id' not in st.session_state:
    st.session_state.user_id = "test_user_001"
user_id = st.session_state.user_id


# --- 辅助函数：大纲覆盖检测 ---
def check_outline_coverage(outline, draft_text):
    if not outline: return []
    coverage = []
    draft_lower = draft_text.lower()
    for point in outline:
        pt_lower = point.lower()
        # 多策略匹配
        is_covered = (
                pt_lower in draft_lower or
                f"{point}：" in draft_text or
                f"{point}:" in draft_text
        )
        coverage.append({"title": point, "covered": is_covered})
    return coverage


# --- 辅助函数：完结检测 ---
def check_if_finished(curr_pos, total_len, outline_coverage):
    # 条件1：物理进度走完
    if curr_pos >= total_len: return True
    # 条件2：大纲覆盖率 > 90%
    if outline_coverage:
        covered_count = sum(1 for item in outline_coverage if item['covered'])
        if covered_count >= len(outline_coverage) * 0.9: return True
    return False

def check_and_update_streak(uid):
    """检查并更新连续打卡天数"""
    try:
        profile = get_user_profile(uid)
        last_date_str = profile.get('last_active_date')
        current_streak = profile.get('study_streak', 0)
        today_str = str(datetime.date.today())

        # 如果今天还没记录
        if last_date_str != today_str:
            new_streak = 1  # 默认重置

            if last_date_str:
                last_date = datetime.datetime.strptime(last_date_str, '%Y-%m-%d').date()
                yesterday = datetime.date.today() - datetime.timedelta(days=1)

                # 如果上次活跃是昨天，天数+1
                if last_date == yesterday:
                    new_streak = current_streak + 1
                # 如果是更早之前，保持为 1 (重置)

            # 更新数据库
            supabase.table("study_profile").update({
                "last_active_date": today_str,
                "study_streak": new_streak
            }).eq("user_id", uid).execute()

            return new_streak
    except Exception as e:
        print(f"Streak Error: {e}")
        return 0

check_and_update_streak(user_id)

# ==============================================================================
# 3. 核心功能函数 (AI / DB / File)
# ==============================================================================

async def _generate_audio_coroutine(text, voice, filepath):
    """内部协程，负责实际生成"""
    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(filepath)

def generate_audio_file(text, voice="zh-CN-XiaoxiaoNeural"):
    """
    [Bug修复] 同步包装器。
    Streamlit 运行时可能已有 Event Loop 或处于特殊线程。
    直接 asyncio.run() 会导致 'There is no current event loop' 错误。
    """
    temp_dir = tempfile.gettempdir()
    filename = f"tts_{uuid.uuid4()}.mp3"
    filepath = os.path.join(temp_dir, filename)

    try:
        # 方案：创建新的事件循环并在其中运行，用完即毁，确保线程安全
        new_loop = asyncio.new_event_loop()
        asyncio.set_event_loop(new_loop)
        new_loop.run_until_complete(_generate_audio_coroutine(text, voice, filepath))
        new_loop.close()
        return filepath
    except Exception as e:
        st.error(f"语音合成失败: {e}")
        return None

# --- 数据库 Helper 函数 ---
def get_user_profile(uid):
    try:
        res = supabase.table("study_profile").select("*").eq("user_id", uid).execute()
        if not res.data:
            supabase.table("study_profile").insert({"user_id": uid}).execute()
            return {}
        return res.data[0]
    except: return {}

def update_settings(uid, settings_dict):
    """更新用户设置"""
    try:
        curr = get_user_profile(uid).get('settings') or {}
        curr.update(settings_dict)
        supabase.table("study_profile").update({"settings": curr}).eq("user_id", uid).execute()
    except: pass




def save_ai_pref():
    """回调：保存模型选择"""
    p = st.session_state.get('ai_provider_select')
    m = None
    if "OpenRouter" in str(p): m = st.session_state.get('or_model_select')
    elif "DeepSeek" in str(p): m = st.session_state.get('ds_model_select')
    elif "Gemini" in str(p): m = st.session_state.get('gl_model_select')
    elif "Glama" in str(p): m = st.session_state.get('glama_model_id')
    if p: update_settings(user_id, {"last_provider": p, "last_used_model": m})

# --- AI 调用 (通用版 + 动态超时) ---
# --- AI 调用 (通用版：支持模型覆盖 + 超时豁免) ---
# --- AI 调用 (通用版：修复模型混淆 Bug + 动态超时) ---
# --- AI 调用 (通用版：支持 Google / DeepSeek / OpenRouter / Glama) ---
# --- AI 调用 (V8.0: 含 Glama 深度调试模式) ---
# --- AI 服务层 (缓存客户端 + 智能重试 + JSON修复) ---

@st.cache_resource
def get_ai_client(provider, api_key, base_url=None):
    """[性能优化] 缓存 AI 客户端连接，避免每次调用都重新握手"""
    # 只有在使用 OpenAI SDK 的时候才初始化 Client
    if "DeepSeek" in provider or "OpenRouter" in provider or "Glama" in provider:
        try:
            return OpenAI(api_key=api_key, base_url=base_url)
        except Exception as e:
            print(f"Client Init Error: {e}")
            return None
    return None


def clean_ai_json(text):
    """[鲁棒性] 清洗 AI 返回的 JSON，去除 Markdown 标记和不合法字符"""
    if not text: return ""
    # 去除 ```json 和 ``` 标记
    text = re.sub(r'```json\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'```', '', text)
    return text.strip()


def call_ai_universal(prompt, history=[], model_override=None, timeout_override=None, max_retries=1):
    """
    [功能增强] 统一 AI 调用入口：支持重试、错误捕获、客户端复用
    """
    # 1. 确定超时设置
    if timeout_override is not None:
        current_timeout = timeout_override
    else:
        profile = get_user_profile(st.session_state.get('user_id', 'test_user'))
        settings = profile.get('settings') or {}
        current_timeout = settings.get('ai_timeout', 60)

    # 2. 确定服务商与模型
    provider = st.session_state.get('selected_provider', 'Gemini')
    target_model = None

    # 优先级：Override > Glama特定 > 通用Session
    if model_override:
        target_model = model_override
    elif "Gemini" in provider:
        target_model = st.session_state.get('google_model_id', 'gemini-1.5-flash')
    elif "DeepSeek" in provider:
        target_model = st.session_state.get('deepseek_model_id', 'deepseek-chat')
    elif "OpenRouter" in provider:
        target_model = st.session_state.get('openrouter_model_id', 'google/gemini-2.0-flash-exp:free')
    elif "Glama" in provider:
        target_model = st.session_state.get('glama_model_id', 'openai/gpt-4o-mini')

    if not target_model: target_model = "gemini-1.5-flash"

    # --- 内部执行函数 (用于重试) ---
    def _execute_call():
        # A. Google Gemini (REST API 模式 - 不依赖 OpenAI SDK)
        if "Gemini" in provider and not model_override:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{target_model}:generateContent?key={API_KEY}"
            headers = {'Content-Type': 'application/json'}
            contents = []
            for h in history:
                role = "user" if h['role'] == 'user' else "model"
                contents.append({"role": role, "parts": [{"text": h['content']}]})
            contents.append({"role": "user", "parts": [{"text": prompt}]})

            resp = requests.post(url, headers=headers, json={"contents": contents}, timeout=current_timeout)
            if resp.status_code == 200:
                return resp.json()['candidates'][0]['content']['parts'][0]['text']
            else:
                raise Exception(f"Gemini API Error {resp.status_code}: {resp.text}")

        # B. OpenAI 兼容模式 (DeepSeek / OpenRouter / Glama)
        else:
            # 准备参数
            api_key = ""
            base_url = ""

            if model_override and "gemini" in model_override and "openrouter" in st.secrets:
                api_key = st.secrets["openrouter"]["api_key"]
                base_url = st.secrets["openrouter"]["base_url"]
            elif "DeepSeek" in provider:
                api_key = st.secrets["deepseek"]["api_key"]
                base_url = st.secrets["deepseek"]["base_url"]
            elif "OpenRouter" in provider:
                api_key = st.secrets["openrouter"]["api_key"]
                base_url = st.secrets["openrouter"]["base_url"]
            elif "Glama" in provider:
                if "glama" in st.secrets:
                    base_url = st.secrets["glama"]["base_url"].strip().rstrip("/")
                    api_key = st.secrets["glama"]["api_key"]
                else:
                    return "❌ Glama Secrets 未配置"

            # 获取或初始化客户端 (利用缓存)
            client = get_ai_client(provider, api_key, base_url)
            if not client: return "AI Client 初始化失败"

            # 构造消息
            messages = [{"role": "system", "content": "你是一位资深会计讲师。回答请使用 Markdown 格式。"}]
            for h in history:
                role = "assistant" if h['role'] == "model" else h['role']
                messages.append({"role": role, "content": h['content']})
            messages.append({"role": "user", "content": prompt})

            # 发起请求
            resp = client.chat.completions.create(
                model=target_model,
                messages=messages,
                temperature=0.7,
                timeout=current_timeout
            )
            return resp.choices[0].message.content

    # --- 重试逻辑 ---
    last_error = ""
    for attempt in range(max_retries + 1):
        try:
            return _execute_call()
        except Exception as e:
            last_error = str(e)
            if attempt < max_retries:
                time.sleep(1)  # 失败后暂停1秒再试
                continue

    return f"❌ AI 调用失败 (已重试{max_retries}次): {last_error}"


def call_ai_json(prompt, model_override=None):
    """
    [新功能] 专门请求 JSON 数据，带自动清洗和解析，防止报错
    """
    # 强制要求 JSON
    json_prompt = prompt + "\n\n请务必只返回纯 JSON 格式，不要包含 ```json 等 Markdown 标记，也不要有多余的解释文字。"

    res = call_ai_universal(json_prompt, model_override=model_override)
    if not res or "Error" in res or "失败" in res:
        return None

    try:
        clean = clean_ai_json(res)
        # 尝试截取第一个 { 到 最后一个 } 或者是 [ 到 ]
        s_obj = clean.find('{');
        e_obj = clean.rfind('}') + 1
        s_arr = clean.find('[');
        e_arr = clean.rfind(']') + 1

        # 智能判断是对象还是数组
        if s_arr != -1 and (s_obj == -1 or s_arr < s_obj):
            return json.loads(clean[s_arr:e_arr])
        elif s_obj != -1:
            return json.loads(clean[s_obj:e_obj])
        else:
            return json.loads(clean)  # 尝试直接解析

    except json.JSONDecodeError:
        print(f"JSON Parse Error. Raw AI Response: {res}")
        return None


# --- 新增：主观题 AI 评分函数 ---
def ai_grade_subjective(user_ans, std_ans, question_content):
    """
    专门用于主观题评分
    返回: {'score': 0-100, 'feedback': '...'}
    """
    if not user_ans or len(user_ans.strip()) < 2:
        return {'score': 0, 'feedback': '未检测到有效作答。'}
        
    prompt = f"""
    【角色】你是一位严谨的会计阅卷老师。
    【任务】请对考生的主观题答案进行评分。
    
    【题目】
    {question_content}
    
    【标准答案】
    {std_ans}
    
    【考生答案】
    {user_ans}
    
    【评分标准】
    1. 满分 100 分。
    2. 核心会计分录、计算结果、关键术语正确即可得分，不纠结文字表述差异。
    3. 如果分录借贷方向反了，直接 0 分。
    4. 如果金额错误但逻辑正确，给 30-50% 分数。
    
    请以纯 JSON 格式返回：
    {{
        "score": 85,
        "feedback": "分录正确，但折旧计算金额有误（应为1000而非1200）。"
    }}
    """
    try:
        # 强制较短超时，避免卡死，评分通常较快
        res = call_ai_universal(prompt, timeout_override=45)
        clean = res.replace("```json","").replace("```","").strip()
        s = clean.find('{'); e = clean.rfind('}')+1
        return json.loads(clean[s:e])
    except Exception as e:
        return {'score': 0, 'feedback': f"AI 阅卷失败: {e}"}


# --- 动态获取模型列表函数 ---
@st.cache_data(ttl=3600)
def fetch_google_models(api_key):
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
        data = requests.get(url, timeout=10).json()
        return [m['name'].replace("models/", "") for m in data.get('models', []) if "generateContent" in m.get('supportedGenerationMethods', [])]
    except: return []

@st.cache_data(ttl=3600)
def fetch_openrouter_models(api_key):
    try:
        url = "https://openrouter.ai/api/v1/models"
        resp = requests.get(url, headers={"Authorization": f"Bearer {api_key}"}, timeout=10)
        if resp.status_code == 200:
            data = resp.json().get('data', [])
            return sorted([
                {'id': m['id'], 'is_free': (float(m.get('pricing',{}).get('prompt',0))==0) or ':free' in m['id']} 
                for m in data
            ], key=lambda x: x['id'])
        return []
    except: return []

# --- 新增：动态获取 Glama 模型列表 ---
@st.cache_data(ttl=3600)
def fetch_glama_models(api_key, base_url):
    """
    从 Glama 获取可用模型列表
    """
    try:
        # 自动修正 Base URL (防止用户填错)
        # Glama 的标准 Base URL 通常是 https://glama.ai/api/gateway/openai/v1
        # 但获取 models 时只需 base_url + /models
        target_url = base_url.rstrip("/") + "/models"
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        resp = requests.get(target_url, headers=headers, timeout=10)
        
        if resp.status_code == 200:
            data = resp.json().get('data', [])
            # 提取模型 ID 并排序
            return sorted([m['id'] for m in data], key=lambda x: x)
        else:
            print(f"Glama Fetch Error: {resp.status_code} - {resp.text}")
            return []
    except Exception as e:
        print(f"Glama Fetch Exception: {e}")
        return []

# --- 数据库 CRUD Helper ---
def get_subjects():
    return supabase.table("subjects").select("*").execute().data

def get_books(sid):
    return supabase.table("books").select("*").eq("subject_id", sid).eq("user_id", user_id).execute().data

def get_chapters(book_id):
    return supabase.table("chapters").select("*").eq("book_id", book_id).order("start_page", desc=False).execute().data

def save_material_v3(chapter_id, content, uid):
    supabase.table("materials").insert({
        "chapter_id": chapter_id, "content": content, "user_id": uid
    }).execute()


def save_questions_v3(q_list, chapter_id, uid, origin="ai"):
    """
    [安全增强版] 替代原有的 save_questions_v3。
    增加了空值校验和错误捕获，防止因为一条数据格式错误导致整个入库失败。
    """
    if not q_list: return

    data_to_insert = []
    timestamp_str = f"Batch-{int(time.time())}"

    for q in q_list:
        # 简单校验：必须有题目内容和答案
        if not q.get('content') or not q.get('correct_answer') or not q.get('question'):
            # 兼容旧逻辑：有的 AI 返回 key 是 question，有的是 content
            content = q.get('content') or q.get('question')
            ans = q.get('correct_answer') or q.get('answer')
            if not content: continue  # 真的没有内容，跳过

            # 修正数据
            q['content'] = content
            q['correct_answer'] = ans

        # 构造数据
        data_to_insert.append({
            "chapter_id": chapter_id,
            "user_id": uid,
            "content": q.get('content') or q.get('question'),
            "options": q.get('options', []),
            "correct_answer": str(q.get('correct_answer') or q.get('answer')),
            "explanation": q.get('explanation', ''),
            "type": q.get('type', 'single'),
            "origin": origin,
            "batch_source": timestamp_str
        })

    if not data_to_insert: return

    try:
        # 执行批量插入
        supabase.table("question_bank").insert(data_to_insert).execute()
    except Exception as e:
        # 记录详细错误日志
        print(f"Database Insert Error: {e}")
        st.error(f"💾 题目入库失败：{e}")


# --- 🆕 新增：测验状态清理函数 (防止缓存中毒) ---
def cleanup_quiz_session():
    """
    清理测验相关的临时数据。
    在【开始新练习】和【退出练习】时调用，确保 Session 干净。
    """
    # 1. 定义需要清理的 Key 前缀
    target_prefixes = (
        'grade_res_',  # AI评分结果
        'sub_state_',  # 题目提交状态锁
        'saved_db_',  # 数据库存库标记
        'q_subj_',  # 主观题输入框内容
        'q_rad_',  # 单选框状态
        'q_',  # 多选框/其他控件状态
        'feedback_',  # AI反馈文本
        'score_'  # 分数
    )

    # 2. 扫描并收集要删除的 Key
    keys_to_remove = [k for k in st.session_state.keys() if k.startswith(target_prefixes)]

    # 3. 执行删除
    for k in keys_to_remove:
        del st.session_state[k]

    # 4. 重置核心控制变量
    # 注意：不要删 'user_id' 等全局配置
    core_keys = ['quiz_active', 'quiz_data', 'q_idx', 'js_start_time']
    for k in core_keys:
        if k in st.session_state:
            del st.session_state[k]

# --- 文件解析 (PDF/Docx) ---
def extract_pdf(file, start=1, end=None, max_pages=50):
    """
    [性能优化版] 替代原有的 extract_pdf。
    保留了函数名，但增加了内存保护、进度条和最大页数限制。
    """
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            total = len(pdf.pages)

            # 自动修正结束页
            if end is None or end > total: end = total
            start = max(1, start)

            # 安全限制：防止用户上传几百页的书直接把内存撑爆
            # 如果后续代码没有传 max_pages，默认限制 50 页
            if (end - start) > max_pages:
                st.warning(f"⚠️ 为保护系统性能，仅读取前 {max_pages} 页 (原请求 {end - start} 页)。")
                end = start + max_pages

            # 进度条 UI (用户能看到进度了)
            progress_bar = st.progress(0)
            status_text = st.empty()

            for i in range(start - 1, end):
                # 显式进度更新
                current_idx = i - (start - 1)
                total_process = end - (start - 1)
                # 防止除以0
                prog_val = (current_idx + 1) / total_process if total_process > 0 else 0
                progress_bar.progress(min(prog_val, 1.0))
                status_text.caption(f"正在读取第 {i + 1} 页...")

                page = pdf.pages[i]

                # 尝试提取表格 (保持原有逻辑，转为 Markdown 表格)
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        row_str = []
                        for row in table:
                            clean_row = [str(cell).replace('\n', ' ') if cell else '' for cell in row]
                            row_str.append("| " + " | ".join(clean_row) + " |")
                        text += "\n".join(row_str) + "\n\n"

                # 提取文本
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i + 1} ---\n{page_text}\n"

                # [关键优化] 没读 10 页清理一次内存，防止 PDF 过大导致页面崩溃
                if i % 10 == 0:
                    gc.collect()

            status_text.empty()
            progress_bar.empty()

        if len(text) < 100:
            st.warning("⚠️ 提取到的文字极少，该 PDF 可能是图片扫描件，AI 无法识别。")

        return text
    except Exception as e:
        st.error(f"PDF 读取出错: {e}")
        return ""


def save_questions_safe(q_list, chapter_id, uid, origin="ai"):
    """
    [数据安全] 批量插入，带错误捕获，不使用不稳定的 transaction 写法
    """
    if not q_list: return

    data_to_insert = []
    for q in q_list:
        # 简单校验
        if not q.get('content') or not q.get('correct_answer'):
            continue

        data_to_insert.append({
            "chapter_id": chapter_id,
            "user_id": uid,
            "content": q['content'],
            "options": q.get('options', []),
            "correct_answer": q['correct_answer'],
            "explanation": q.get('explanation', ''),
            "type": q.get('type', 'single'),
            "origin": origin,
            "batch_source": f"Batch-{int(time.time())}"
        })

    if not data_to_insert: return

    try:
        # Supabase Python SDK 的 insert 通常是原子的 (单次 HTTP 请求)
        res = supabase.table("question_bank").insert(data_to_insert).execute()
        return res
    except Exception as e:
        # 记录详细错误日志
        print(f"Database Insert Error: {e}")
        st.error("💾 题目入库失败，请检查网络或数据格式。")
        return None

def extract_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except: return ""


# --- 🎓 AI 课堂专用辅助函数 (修复版) ---

@st.cache_data(show_spinner=False)
def get_cached_outline_v2(chapter_id, text_content, uid):
    """
    [安全增强版] 用户隔离的大纲缓存
    通过组合用户ID+章节ID+内容哈希作为缓存键，确保多用户隔离，防止串台。
    """
    # 1. 生成唯一标识符 (Content Hash)
    content_hash = hashlib.md5(text_content[:5000].encode('utf-8')).hexdigest()[:8]

    # 2. 截取首尾中三段作为摘要，减少 Token 消耗
    summary_context = text_content[:3000] + "\n...\n" + text_content[
        len(text_content) // 2: len(text_content) // 2 + 2000]

    prompt = f"""
    【任务】快速扫描教材，列出本章 5-8 个核心知识点标题。
    【教材片段】{summary_context}
    【格式】请返回纯 JSON 字符串数组，例如：["总论", "存货的初始计量", "期末计量"]
    【注意】标题要简洁，不要带序号。
    """
    try:
        # 复用全局定义的 call_ai_json
        res = call_ai_json(prompt)
        if isinstance(res, list) and len(res) > 0:
            return res
        return ["本章概览", "核心考点", "实务案例", "章节总结"]  # 兜底
    except:
        return ["本章概览", "核心考点", "实务案例", "章节总结"]


def check_outline_coverage_v2(outline, draft_text):
    """
    [修复] 更精准的覆盖率检测 (支持 Markdown 标题识别)
    """
    if not outline: return []
    coverage = []
    draft_lower = draft_text.lower()

    for point in outline:
        pt_lower = point.lower()

        # 策略1: 显式标题检测 (## 标题)
        is_header = f"## {point}" in draft_text or f"### {point}" in draft_text or f"#### {point}" in draft_text

        # 策略2: 关键词存在且周围有足够文本
        idx = draft_lower.find(pt_lower)
        has_context = False
        if idx != -1:
            # 检查关键词后是否有至少 50 个字符的内容
            context_chunk = draft_text[idx + len(point): idx + len(point) + 100]
            if len(context_chunk.strip()) > 30:
                has_context = True

        is_covered = is_header or has_context
        coverage.append({"title": point, "covered": is_covered})
    return coverage


def check_if_finished_v2(curr_pos, total_len, outline_coverage):
    """
    [修复] 综合完结检测
    """
    # 1. 物理进度 (达到 100%)
    if curr_pos >= total_len: return True

    # 2. 大纲覆盖率 (严格模式 > 90%)
    if outline_coverage:
        covered_count = sum(1 for item in outline_coverage if item['covered'])
        total_pts = len(outline_coverage)
        if total_pts > 0 and (covered_count / total_pts) >= 0.9:
            return True

    return False

# ==============================================================================
# 4. 侧边栏与导航 (修复版：统一菜单名称)
# ==============================================================================
profile = get_user_profile(user_id)
settings = profile.get('settings') or {}

with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3135/3135715.png", width=50)
    st.markdown("### 会计私教 Pro")
    
    # --- AI 设置 (保持不变) ---
    provs = ["Gemini (官方直连)", "DeepSeek (官方直连)", "OpenRouter (聚合平台)", "Glama (聚合平台)"] # <--- 加在这里
    saved_p = settings.get('last_provider')
    idx_p = 0
    if saved_p:
        for i, x in enumerate(provs):
            if saved_p in x: idx_p = i; break
            
    prov = st.selectbox("🧠 AI 大脑", provs, index=idx_p, key="ai_provider_select", on_change=save_ai_pref)
    st.session_state.selected_provider = prov
    
    saved_m = settings.get('last_used_model')
    
    # 1. Gemini
    if "Gemini" in prov:
        opts = fetch_google_models(st.secrets["GOOGLE_API_KEY"]) or ["gemini-1.5-flash"]
        idx_m = opts.index(saved_m) if saved_m in opts else 0
        st.session_state.google_model_id = st.selectbox("🔌 模型", opts, index=idx_m, key="gl_model_select", on_change=save_ai_pref)
        
    # 2. DeepSeek
    elif "DeepSeek" in prov:
        opts = ["deepseek-chat", "deepseek-reasoner"]
        idx_m = opts.index(saved_m) if saved_m in opts else 0
        st.session_state.deepseek_model_id = st.selectbox("🔌 模型", opts, index=idx_m, key="ds_model_select", on_change=save_ai_pref)
        
    # 3. OpenRouter
    elif "OpenRouter" in prov:
        # (保持原有的 OpenRouter 代码不变...)
        try:
            all_ms = fetch_openrouter_models(st.secrets["openrouter"]["api_key"])
            if not all_ms: final_ids = ["google/gemini-2.0-flash-exp:free"]
            else:
                ft = st.radio("筛选", ["🤑 免费", "🌎 全部"], horizontal=True)
                subset = [m for m in all_ms if m['is_free']] if "免费" in ft else all_ms
                final_ids = [m['id'] for m in subset]
                if not final_ids: final_ids = [m['id'] for m in all_ms]
            idx_m = final_ids.index(saved_m) if saved_m in final_ids else 0
            st.session_state.openrouter_model_id = st.selectbox("🔌 模型", final_ids, index=idx_m, key="or_model_select", on_change=save_ai_pref)
        except: st.error("OpenRouter 连接失败")

    # 4. Glama (自动获取模型版)
# -------------------------------------------------------
    # Glama (稳定版：预设列表 + 手动输入)
    # -------------------------------------------------------
    elif "Glama" in prov:
        st.caption("🚀 已启用 Glama 网关加速")
        
        # 1. 定义 Glama 支持的常用模型 (根据官方文档整理)
        glama_presets = [
            "google-vertex/gemini-2.0-flash-exp",  # 👈 加上 google-vertex/ 前缀
            "google-vertex/gemini-1.5-pro",
            "openai/gpt-4o",
            "openai/gpt-4o-mini",
            "anthropic/claude-3-5-sonnet",
            "meta-llama/llama-3.1-405b-instruct"
        ]
        
        # 2. 提供切换方式：选常用的 vs 手输冷门的
        input_mode = st.radio("模型选择", ["⚡ 常用模型", "⌨️ 手动输入"], horizontal=True, label_visibility="collapsed")
        
        if "常用" in input_mode:
            # 自动定位上次选的模型
            idx_m = glama_presets.index(saved_m) if saved_m in glama_presets else 0
            st.session_state.glama_model_id = st.selectbox(
                "🔌 选择模型", 
                glama_presets, 
                index=idx_m, 
                key="glama_list_select", 
                on_change=save_ai_pref
            )
        else:
            st.session_state.glama_model_id = st.text_input(
                "请输入模型 ID", 
                value=saved_m or "gemini-2.0-flash-exp", 
                placeholder="例如: google-vertex/gemini-1.5-flash",
                key="glama_manual_input",
                on_change=save_ai_pref
            )
            st.caption("提示：可在 Glama 后台查看完整的 Model ID")

    
    # --- 导航菜单 (关键修改点：名字与下方主逻辑严格一致) ---
    # 定义菜单列表
    MENU_OPTIONS = [
        "🏠 仪表盘",
        "📂 智能拆书 & 资料",
        "🎓 AI 课堂 (讲义)",
        "📝 章节特训",    # 注意：这里去掉了"(刷题)"
        "⚔️ 全真模考",
        "📊 弱项分析",
        "❌ 错题本",
        "🛠️ 数据管理 & 补录",
        "⚙️ 设置中心"
    ]
    
    menu = st.radio("功能导航", MENU_OPTIONS, label_visibility="collapsed")
    
    # --- 倒计时 ---
    if profile.get('exam_date'):
        try:
            target = datetime.datetime.strptime(profile['exam_date'], '%Y-%m-%d').date()
            today = datetime.date.today()

            # 核心修复逻辑：
            if target < today:
                # 如果存的日期过期了，先看看“今年的考试”是不是还没到？
                this_year_exam = datetime.date(today.year, 9, 6)  # 假设考试在9月

                if today < this_year_exam:
                    # 如果还没到今年的9月，目标就是今年
                    target = this_year_exam
                    label = f"{today.year}赛季"
                else:
                    # 如果今年9月也过了，那就备战明年
                    target = datetime.date(today.year + 1, 9, 6)
                    label = f"{today.year + 1}赛季"

                days = (target - today).days
                st.metric("⏳ 备战考试", f"{days} 天", delta=label)
            else:
                days = (target - today).days
                st.metric("⏳ 距离考试", f"{days} 天", delta="冲刺" if days < 30 else "稳住")
        except Exception as e:
            print(f"Date Error: {e}")

# ==============================================================================
# 5. 各页面主逻辑 (V3.0 完整复刻版)
# ==============================================================================

# === 🏠 仪表盘 (Bento Grid 风格) ===
if menu == "🏠 仪表盘":
    # 1. 欢迎语与智能倒计时
    exam_date_str = profile.get('exam_date')
    today = datetime.date.today()
    days_left = 0
    display_year = today.year

    if exam_date_str:
        try:
            target_date = datetime.datetime.strptime(exam_date_str, '%Y-%m-%d').date()

            # 核心修复逻辑：
            if target_date < today:
                # 检查今年考试是否过期
                this_year_exam = datetime.date(today.year, 9, 6)
                if today < this_year_exam:
                    target_date = this_year_exam
                    display_year = today.year
                else:
                    target_date = datetime.date(today.year + 1, 9, 6)
                    display_year = today.year + 1
            else:
                display_year = target_date.year

            days_left = (target_date - today).days
        except:
            days_left = 0

    # 动态标题
    title_html = f"### 🍂 备战 <span style='color:#00C090'>{display_year}</span>"
    msg = "种一棵树最好的时间是十年前，其次是现在。"

    # 如果天数很少（比如冲刺阶段），换个标语
    if days_left > 0 and days_left < 60:
        title_html = f"### 🌞 冲刺 <span style='color:#ff4b4b'>{days_left}</span> 天"
        msg = "现在的从容，就是考场上的噩梦。"

    st.markdown(title_html, unsafe_allow_html=True)
    st.info(f"👨‍🏫 **班主任说：** {msg}")

    # 2. 核心数据 Bento Grid
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""
        <div class="css-card">
            <i class="bi bi-collection-fill stat-icon"></i>
            <div class="stat-title">累计刷题</div>
            <div class="stat-value">{profile.get('total_questions_done', 0)}</div>
        </div>
        """, unsafe_allow_html=True)
    with c2:
        st.markdown(f"""
        <div class="css-card">
            <i class="bi bi-fire stat-icon" style="color:#FF7043"></i>
            <div class="stat-title">连续打卡</div>
            <div class="stat-value">{profile.get('study_streak', 0)} <span style="font-size:1rem">天</span></div>
        </div>
        """, unsafe_allow_html=True)
    with c3:
        # 简单计算错题数
        try:
            err_count = supabase.table("user_answers").select("id", count="exact").eq("user_id", user_id).eq("is_correct", False).execute().count
        except: err_count = 0
        st.markdown(f"""
        <div class="css-card">
            <i class="bi bi-bookmark-x-fill stat-icon" style="color:#dc3545"></i>
            <div class="stat-title">待消灭错题</div>
            <div class="stat-value">{err_count}</div>
        </div>
        """, unsafe_allow_html=True)

# =========================================================
# 📂 智能拆书 & 资料 (V8.6: 修复教材模式死锁与Prompt适配)
# =========================================================
elif menu == "📂 智能拆书 & 资料":
    st.title("📂 资料库管理 (Pro)")


    # --- 0. 辅助函数定义 (置顶防止 NameError) ---

    def clean_textbook_content(text):
        """
        [增强版] 教材文本清洗
        修复乱码、去除控制字符、标准化 Unicode
        """
        if not text: return ""

        # 1. Unicode 标准化 (NFKC 模式)
        # 这步非常关键！它会把兼容字符（如合字 ﬁ）拆分为标准字符 (fi)
        # 也会修复很多看起来像乱码的拉丁字符
        text = unicodedata.normalize('NFKC', text)

        # 2. 清除不可见控制字符 (除了换行符 \n 和制表符 \t)
        # \x00-\x08: Null等
        # \x0b-\x0c: 垂直制表等
        # \x0e-\x1f: 其他控制符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        # 3. 替换特殊的空白符号为普通空格
        text = text.replace('\xa0', ' ')  # No-break space
        text = text.replace('\u3000', ' ')  # 全角空格

        # 4. 针对性修复常见的 PDF 乱码 (根据你的截图定制)
        # 很多 PDF 会把空格识别成 â 或 ð 等，这里可以手动暴力替换
        # 如果你发现特定的怪符号总是出现，可以在这里加
        text = text.replace('â', '')
        text = text.replace('ð', '')

        # 5. 原有的按行清洗逻辑
        lines = text.split('\n')
        cleaned = []
        for line in lines:
            s = line.strip()
            # 过滤掉纯数字(页码)、过短的噪音
            # 增加逻辑：如果一行全是乱七八糟的符号（非中英文数字），也丢弃
            if len(s) < 2 or (s.isdigit() and len(s) < 5): continue
            cleaned.append(s)

        return "\n".join(cleaned)


    subjects = get_subjects()
    if not subjects: st.error("请先初始化科目数据"); st.stop()

    # 1. 顶层选择区
    c1, c2 = st.columns([1, 2])
    with c1:
        s_name = st.selectbox("1. 所属科目", [s['name'] for s in subjects])
        sid = next(s['id'] for s in subjects if s['name'] == s_name)
    with c2:
        books = get_books(sid)
        book_map = {f"{b['title']} (ID:{b['id']})": b['id'] for b in books}
        b_opts = ["➕ 上传新资料...", "---"] + list(book_map.keys())
        sel_book_label = st.selectbox("2. 选择书籍/文件", b_opts)

    st.divider()

    # =====================================================
    # 场景 A: 上传新资料
    # =====================================================
    if "➕ 上传新资料" in sel_book_label:
        st.markdown("#### 📤 第一步：选择导入方式")

        import_mode = st.radio("请选择资料类型",
                               ["📄 PDF 智能拆分 (适合整本扫描书/习题册)",
                                "📊 Excel/CSV 结构化导入 (适合已整理的讲义/考点/法条)"],
                               horizontal=True
                               )

        # -------------------------------------------------
        # 模式 1: PDF 智能拆分
        # -------------------------------------------------
        if "PDF" in import_mode:
            st.info("💡 AI 将自动分析目录结构，并提取题目或正文。")

            doc_type = st.radio("文件内容是？", ["📑 习题库 (录入题目)", "📖 纯教材 (AI导学)"], horizontal=True)
            up_file = st.file_uploader("拖入 PDF 文件", type="pdf")

            if up_file:
                try:
                    # 预读取页数
                    with pdfplumber.open(up_file) as pdf:
                        total_pages = len(pdf.pages)

                    # 初始化配置状态
                    if 'toc_config' not in st.session_state:
                        st.session_state.toc_config = {
                            "toc_s": 1, "toc_e": min(10, total_pages),
                            "content_s": 1
                        }

                    # --- Step 1: 目录结构分析配置 ---
                    if 'toc_result' not in st.session_state:
                        c_info = st.container()
                        with c_info:
                            st.markdown(f"✅ 文件已加载，共 **{total_pages}** 页。")

                            # 无目录模式开关
                            col_toc_flag, col_dummy = st.columns([2, 3])
                            with col_toc_flag:
                                is_no_toc = st.checkbox("🚫 本文档无目录 (视为单章节或手动分节)", value=False,
                                                        help="勾选后将跳过AI目录分析，直接建立一个包含全书的章节。")

                            st.divider()
                            col_toc, col_body = st.columns(2)

                            # 1. 目录设置区
                            with col_toc:
                                if not is_no_toc:
                                    st.markdown("**1. 目录范围**")
                                    ts = st.number_input("目录开始页", 1, total_pages,
                                                         st.session_state.toc_config['toc_s'])
                                    te = st.number_input("目录结束页", 1, total_pages,
                                                         st.session_state.toc_config['toc_e'])
                                else:
                                    st.markdown("**1. 目录范围**")
                                    st.info("已禁用目录分析")
                                    ts, te = 1, 1

                                    # 2. 正文设置区
                            with col_body:
                                st.markdown("**2. 正文起始**")
                                cs = st.number_input("正文(第一章/内容)开始页", 1, total_pages,
                                                     st.session_state.toc_config['content_s'])

                            # 3. 答案位置 (仅习题库)
                            ans_mode = "无"
                            as_page = 0
                            if "习题库" in doc_type:
                                st.markdown("**3. 答案位置**")
                                ans_mode = st.radio("答案在哪？",
                                                    ["🅰️ 紧跟在题目后面", "🅱️ 集中在文件末尾", "🇨 集中在每一章末尾"])
                                if ans_mode == "🅱️ 集中在文件末尾":
                                    as_page = st.number_input("答案区域开始页", 1, total_pages,
                                                              value=max(1, total_pages - 5))

                            # --- Prompt 控制区 ---
                            user_toc_prompt = ""
                            if not is_no_toc:
                                st.markdown("---")
                                with st.expander("🛠️ AI 指令微调 (目录分析)", expanded=False):
                                    # 🟢 关键修复：根据文档类型切换 Prompt
                                    if "纯教材" in doc_type:
                                        default_toc_prompt = f"""
任务：分析教材目录，提取章节结构。
总页数：{total_pages}。
正文起始偏移：正文内容始于第 {cs} 页。

请提取“章”或“节”的名称，并推算其在PDF的【物理起始页码】。
忽略前言、附录。

返回格式示例：[{{ "title": "第一章 总论", "start_page": 5, "end_page": 20 }}]
                                        """
                                    else:
                                        default_toc_prompt = f"""
任务：分析习题册目录。
总页数：{total_pages}。
正文起始：第 {cs} 页。
请提取章节名称，推算物理页码。
返回格式：[{{ "title": "第一章 存货", "start_page": 5, "end_page": 10 }}]
                                        """
                                    user_toc_prompt = st.text_area("提示词", value=default_toc_prompt.strip(),
                                                                   height=150)

                            # --- 执行按钮 ---
                            btn_label = "🚀 生成全书结构" if is_no_toc else "🚀 执行AI目录分析"

                            if st.button(btn_label, type="primary"):
                                st.toast("正在处理中...")
                                # A. 无目录模式
                                if is_no_toc:
                                    mock_data = [{
                                        "title": "全文内容 (自动生成)",
                                        "start_page": cs,
                                        "end_page": total_pages,
                                        "ans_start_page": as_page if "文件末尾" in ans_mode else 0,
                                        "ans_end_page": total_pages if "文件末尾" in ans_mode else 0
                                    }]
                                    st.session_state.toc_result = mock_data
                                    st.session_state.ans_mode_cache = ans_mode
                                    st.rerun()

                                # B. AI 分析模式
                                else:
                                    with st.spinner("AI 正在阅读目录，请稍候..."):
                                        try:
                                            up_file.seek(0)
                                            toc_txt = extract_pdf(up_file, ts, te)
                                            if not toc_txt.strip():
                                                st.error("⚠️ 未能从指定页码提取到文字，可能是图片扫描件？")
                                            else:
                                                full_p = f"{user_toc_prompt}\n\n目录文本：\n{toc_txt[:10000]}"
                                                res = call_ai_universal(full_p)

                                                if res and "Error" not in res:
                                                    clean = res.replace("```json", "").replace("```", "").strip()
                                                    s = clean.find('[');
                                                    e = clean.rfind(']') + 1
                                                    data = json.loads(clean[s:e])

                                                    # 补全字段
                                                    for row in data:
                                                        row['ans_start_page'] = as_page if "文件末尾" in ans_mode else 0
                                                        row[
                                                            'ans_end_page'] = total_pages if "文件末尾" in ans_mode else 0

                                                    st.session_state.toc_result = data
                                                    st.session_state.ans_mode_cache = ans_mode
                                                    st.rerun()
                                                else:
                                                    st.error(f"AI 响应异常: {res}")
                                        except Exception as e:
                                            st.error(f"分析出错: {e}")

                    # --- Step 2: 确认结构 ---
                    if 'toc_result' in st.session_state:
                        st.divider()
                        c_head, c_re = st.columns([4, 1])
                        with c_head:
                            st.markdown("#### 📝 第二步：确认章节结构")
                        with c_re:
                            if st.button("🔄 重做第一步"):
                                del st.session_state.toc_result
                                st.rerun()

                        cached_ans_mode = st.session_state.get('ans_mode_cache', '无')
                        is_textbook = "纯教材" in doc_type

                        lbl_start = "正文起始页" if is_textbook else "题目起始页"
                        lbl_end = "正文结束页" if is_textbook else "题目结束页"

                        col_cfg = {
                            "title": "章节名称",
                            "start_page": st.column_config.NumberColumn(lbl_start, format="%d", min_value=1),
                            "end_page": st.column_config.NumberColumn(lbl_end, format="%d", min_value=1)
                        }

                        if not is_textbook and "文件末尾" in cached_ans_mode:
                            col_cfg["ans_start_page"] = st.column_config.NumberColumn("答案起始", format="%d")
                            col_cfg["ans_end_page"] = st.column_config.NumberColumn("答案结束", format="%d")

                        edited_df = st.data_editor(st.session_state.toc_result, column_config=col_cfg,
                                                   num_rows="dynamic", use_container_width=True)

                        # --- Step 3: 提取与入库 ---

                        # >>> 分支 A: 习题库 (含判断题支持 & 跨页修复) <<<
                        if "习题库" in doc_type:
                            st.divider()
                            st.markdown("#### 🧪 第三步：入库配置与测试")

                            st.info("💡 如果题目不全或答案丢失，请增大【跨页缓冲】。")
                            page_buffer = st.slider("📐 跨页缓冲 (自动多读N页)", 0, 5, 1,
                                                    help="防止答案刚好在下一页被截断。")

                            st.markdown("🛠️ **AI 指令微调 (题目提取)**")

                            # 终极融合版 Prompt
                            default_extract_prompt = """
【角色】你是一个专业的文档数据清洗专家。
【任务】从包含噪音（页码、页眉、换行符）的 PDF 文本中提取题目。

【核心处理规则】
1. 🧹 **跨页与噪音修复（至关重要）**：
   - 文本中可能包含 "Page 10", "2024年真题" 等页眉页脚噪音，请直接忽略。
   - **如果一句话在行末中断（跨页），请务必将其与下一行拼接，还原为完整句子。**
   - 不要把因为排版原因断开的题目拆成两道题。

2. 🧠 **题型分类与清洗**：
   - **判断题 (Judgment)**：
     特征：题目是陈述句，要求判断对错（√/×, T/F）。
     处理：type="judgment"。**answer 请统一转为 "A"(对) 或 "B"(错)。**

   - **单/多选题 (Single/Multi)**：
     特征：带有 A,B,C,D 选项。
     处理：必须提取完整选项。如果原文选项没有 ABCD 编号，请自动补全编号。

   - **主观题 (Subjective)**：
     特征：简答、计算分析、综合题、分录题。
     **处理关键**：务必将“背景资料”（如甲公司2024发生如下业务...）与“所有小问的要求”**合并**存入 question 字段。
     **严禁**将一道大题的 (1)(2)(3) 小问拆成多条记录，必须合并为一条 type="subjective" 的数据。

【返回格式】
纯 JSON 列表，无 Markdown 标记：
[
  {
    "question": "1. [判断] 企业应当按月计提折旧。",
    "type": "judgment",
    "options": [], 
    "answer": "A", 
    "explanation": "..."
  },
  {
    "question": "【计算题】甲公司资料如下... (1)计算净利润; (2)做分录...",
    "type": "subjective",
    "options": [],
    "answer": "(1) 100万; (2) 借:...", 
    "explanation": "..."
  }
]
                            """
                            user_extract_prompt = st.text_area("提取提示词", value=default_extract_prompt.strip(),
                                                               height=250)

                            # 预览功能
                            preview_idx = st.selectbox("选择章节测试", range(len(edited_df)),
                                                       format_func=lambda x: edited_df[x]['title'])

                            if st.button("🔍 抽取 5 题测试"):
                                row = edited_df[preview_idx]
                                try:
                                    # 提取题目文本
                                    p_s = int(float(row['start_page']))
                                    p_e = min(p_s + 3, int(float(row['end_page'])))
                                    up_file.seek(0)
                                    q_text = extract_pdf(up_file, p_s, p_e)

                                    # 提取答案文本
                                    if "文件末尾" in cached_ans_mode:
                                        a_s = int(float(row['ans_start_page']))
                                        a_e = min(a_s + 3 + page_buffer, int(float(row['ans_end_page'])))
                                        up_file.seek(0)
                                        a_text = extract_pdf(up_file, a_s, a_e)
                                        q_text += f"\n\n====== 答案区域 (缓冲 {page_buffer} 页) ======\n{a_text}"

                                    full_p = f"{user_extract_prompt}\n\n待提取文本：\n{q_text[:25000]}"

                                    with st.spinner("AI 正在提取..."):
                                        res = call_ai_universal(full_p)
                                        if "QuotaFailure" in str(res):
                                            st.error("⚠️ API 配额超限。")
                                        elif res:
                                            cln = res.replace("```json", "").replace("```", "").strip()
                                            s = cln.find('[');
                                            e = cln.rfind(']') + 1
                                            st.session_state.preview_data = json.loads(cln[s:e])
                                except Exception as e:
                                    st.error(f"测试失败: {e}")

                            # 展示结果
                            if st.session_state.get('preview_data'):
                                st.write("##### 👀 识别结果预览")
                                p_df = pd.DataFrame(st.session_state.preview_data)
                                st.dataframe(p_df[['type', 'question', 'answer']], use_container_width=True)

                                # 执行全量
                                if st.button("💾 确认无误，执行全量入库", type="primary"):
                                    progress_bar = st.progress(0)
                                    st_text = st.empty()

                                    # 1. 建书
                                    b_res = supabase.table("books").insert({
                                        "user_id": user_id, "subject_id": sid,
                                        "title": up_file.name.replace(".pdf", ""), "total_pages": total_pages
                                    }).execute()
                                    bid = b_res.data[0]['id']

                                    try:
                                        for i, row in enumerate(edited_df):
                                            st_text.text(f"正在处理：{row['title']}...")
                                            c_s = int(float(row['start_page']));
                                            c_e = int(float(row['end_page']))

                                            c_res = supabase.table("chapters").insert({
                                                "book_id": bid, "title": row['title'], "start_page": c_s,
                                                "end_page": c_e, "user_id": user_id
                                            }).execute()
                                            cid = c_res.data[0]['id']

                                            # 提取内容
                                            up_file.seek(0)
                                            txt = extract_pdf(up_file, c_s, c_e)

                                            if "文件末尾" in cached_ans_mode:
                                                a_s = int(float(row['ans_start_page']))
                                                a_e_original = int(float(row['ans_end_page']))
                                                a_e_safe = min(a_e_original + page_buffer, total_pages)
                                                if a_s > 0:
                                                    up_file.seek(0)
                                                    a_text = extract_pdf(up_file, a_s, a_e_safe)
                                                    txt += f"\n\n====== 答案区域 ======\n{a_text}"

                                            # 调用 AI
                                            final_p = f"{user_extract_prompt}\n\n文本：\n{txt[:60000]}"
                                            r = call_ai_universal(final_p, timeout_override=300)

                                            if r and "QuotaFailure" not in str(r):
                                                try:
                                                    cln = r.replace("```json", "").replace("```", "").strip()
                                                    s = cln.find('[');
                                                    e = cln.rfind(']') + 1
                                                    qs = json.loads(cln[s:e])

                                                    db_data = []
                                                    for q in qs:
                                                        # === 入库清洗逻辑 ===
                                                        raw_type = q.get('type', 'single').lower()
                                                        final_type = 'single'
                                                        final_opts = q.get('options', [])
                                                        final_ans = str(q.get('answer', '')).strip().upper()

                                                        if 'judgment' in raw_type or '判断' in raw_type:
                                                            final_type = 'judgment'
                                                            if not final_opts: final_opts = ["A. 正确", "B. 错误"]
                                                            if final_ans in ['T', 'TRUE', '√', '正确', '对']:
                                                                final_ans = 'A'
                                                            elif final_ans in ['F', 'FALSE', '×', '错误', '错']:
                                                                final_ans = 'B'

                                                        elif 'subjective' in raw_type or not final_opts or len(
                                                                final_ans) > 10:
                                                            final_type = 'subjective'
                                                        elif len(final_ans) > 1 or 'multi' in raw_type:
                                                            final_type = 'multi'

                                                        db_data.append({
                                                            "chapter_id": cid, "user_id": user_id,
                                                            "content": q['question'],
                                                            "options": final_opts,
                                                            "correct_answer": final_ans,
                                                            "explanation": q.get('explanation', ''),
                                                            "type": final_type,
                                                            "origin": "extract",
                                                            "batch_source": "PDF-V8.5"
                                                        })
                                                    if db_data:
                                                        supabase.table("question_bank").insert(db_data).execute()
                                                except:
                                                    pass

                                            progress_bar.progress((i + 1) / len(edited_df))

                                        progress_bar.progress(100)
                                        st.balloons()
                                        st.success(f"🎉 入库完成！书籍《{up_file.name}》已保存。")

                                        st.markdown("---")
                                        if st.button("🔄 继续上传新资料", type="primary", key="btn_continue_pdf"):
                                            keys_to_clear = ['toc_result', 'toc_config', 'preview_data',
                                                             'ans_mode_cache']
                                            for k in keys_to_clear:
                                                if k in st.session_state: del st.session_state[k]
                                            st.rerun()
                                    except Exception as e:
                                        st.error(f"出错: {e}")

                        # >>> 分支 B: 纯教材 (核心修复部分) <<<
                        elif "纯教材" in doc_type:
                            st.divider()
                            st.markdown("#### 💾 第三步：执行教材入库")
                            st.info("系统将按章节切割 PDF，提取纯文本并存入【Materials】表，供 AI 课堂调用。")

                            if st.button("🚀 开始导入教材", type="primary"):
                                try:
                                    b_res = supabase.table("books").insert({
                                        "user_id": user_id, "subject_id": sid,
                                        "title": up_file.name.replace(".pdf", ""), "total_pages": total_pages
                                    }).execute()
                                    bid = b_res.data[0]['id']

                                    bar = st.progress(0)
                                    status_txt = st.empty()

                                    for i, row in enumerate(edited_df):
                                        chap_title = row['title']
                                        status_txt.text(f"正在处理：{chap_title} ...")

                                        c_s = int(float(row['start_page']))
                                        c_e = int(float(row['end_page']))

                                        c_res = supabase.table("chapters").insert({
                                            "book_id": bid, "title": chap_title,
                                            "start_page": c_s, "end_page": c_e, "user_id": user_id
                                        }).execute()

                                        up_file.seek(0)
                                        txt = extract_pdf(up_file, c_s, c_e)
                                        clean_txt = clean_textbook_content(txt)
                                        if clean_txt:
                                            save_material_v3(c_res.data[0]['id'], clean_txt, user_id)

                                        bar.progress((i + 1) / len(edited_df))

                                    bar.progress(100)
                                    st.balloons()
                                    st.success(f"🎉 教材《{up_file.name}》入库成功！")

                                    st.markdown("---")
                                    if st.button("🔄 继续上传", type="primary"):
                                        del st.session_state.toc_result
                                        st.rerun()
                                except Exception as e:
                                    st.error(f"导入失败: {e}")

                except Exception as e:
                    st.error(f"文件预读错误: {e}")

        # -------------------------------------------------
        # 模式 2: Excel 结构化导入
        # -------------------------------------------------
        else:
            st.markdown("#### 📥 Excel 教材导入")
            st.info("💡 适合导入已整理好的笔记、考点汇总、法条大全。**无需消耗 AI Token，内容 100% 准确。**")

            data_template = [
                {"章节名称": "第一章 总论", "正文内容": "这里填入第一章的所有知识点文本..."},
                {"章节名称": "第二章 存货", "正文内容": "存货的初始计量包括：\n1. 购买价款..."}
            ]
            df_temp = pd.DataFrame(data_template)
            csv = df_temp.to_csv(index=False).encode('utf-8-sig')
            st.download_button("⬇️ 下载导入模版 (.csv)", csv, "教材导入模版.csv", "text/csv")

            st.divider()

            up_excel = st.file_uploader("上传填好的文件", type=["csv", "xlsx"])
            book_name_input = st.text_input("给这份资料起个名字", placeholder="例如：2025中级实务-考点狂背版")

            if up_excel and book_name_input:
                if st.button("🚀 立即导入数据库", type="primary"):
                    try:
                        if up_excel.name.endswith('.csv'):
                            df = pd.read_csv(up_excel)
                        else:
                            df = pd.read_excel(up_excel)

                        bar = st.progress(0)

                        b_res = supabase.table("books").insert({
                            "user_id": user_id, "subject_id": sid, "title": book_name_input, "total_pages": 0
                        }).execute()
                        bid = b_res.data[0]['id']

                        total_rows = len(df)
                        for i, row in df.iterrows():
                            chap_title = str(row.get('章节名称') or row.get('title') or f'第 {i + 1} 节').strip()
                            content = str(row.get('正文内容') or row.get('content') or '').strip()
                            if not content: continue

                            c_res = supabase.table("chapters").insert({
                                "book_id": bid, "title": chap_title, "start_page": 0, "end_page": 0, "user_id": user_id
                            }).execute()
                            cid = c_res.data[0]['id']
                            save_material_v3(cid, content, user_id)
                            bar.progress((i + 1) / total_rows)

                        bar.progress(100)
                        st.balloons()
                        st.success(f"🎉 导入成功！已创建书籍：《{book_name_input}》")

                        st.markdown("---")
                        if st.button("🔄 继续导入下一个 Excel", type="primary", key="btn_continue_excel"):
                            st.rerun()

                    except Exception as e:
                        st.error(f"导入失败: {e}。\n请确保 Excel 包含【章节名称】和【正文内容】两列。")

    # =====================================================
    # 场景 B: 已有书籍管理
    # =====================================================
    elif books:
        if sel_book_label == "---":
            st.info("👈 请选择一本书籍进行管理")
        else:
            bid = book_map[sel_book_label]
            curr_book_info = next((b for b in books if b['id'] == bid), {})

            # 书籍头部信息
            c_tit, c_act = st.columns([5, 1])
            with c_tit:
                st.markdown(f"### 📘 {curr_book_info.get('title', '未知书籍')}")
            with c_act:
                if st.button("🗑️ 删除本书", type="primary"):
                    try:
                        supabase.table("books").delete().eq("id", bid).execute()
                        st.toast("书籍已删除")
                        time.sleep(1)
                        st.rerun()
                    except:
                        st.error("删除失败")

            # 书籍重命名/转科设置
            with st.expander("🔧 书籍设置 (修正科目 / 重命名)", expanded=False):
                c_set1, c_set2, c_set3 = st.columns([2, 2, 1])
                with c_set1:
                    new_title = st.text_input("📖 书籍名称", value=curr_book_info.get('title', ''))
                with c_set2:
                    all_subs = get_subjects()
                    all_sub_names = [s['name'] for s in all_subs]
                    # 防止索引越界
                    curr_sub_idx = 0
                    if s_name in all_sub_names:
                        curr_sub_idx = all_sub_names.index(s_name)
                    target_sub_name = st.selectbox("🔀 归属科目", all_sub_names, index=curr_sub_idx)
                with c_set3:
                    st.write("");
                    st.write("")
                    if st.button("💾 保存变更"):
                        try:
                            target_sid = next(s['id'] for s in all_subs if s['name'] == target_sub_name)
                            supabase.table("books").update({
                                "title": new_title, "subject_id": target_sid
                            }).eq("id", bid).execute()
                            st.success("✅ 更新成功！")
                            time.sleep(1)
                            st.rerun()
                        except Exception as e:
                            st.error(f"修改失败: {e}")

            st.divider()

            # === 章节列表与内容预览 ===
            chapters = get_chapters(bid)
            if not chapters:
                st.info("本书暂无章节，请去上方重新拆分或导入。")
            else:
                st.write(f"📚 共找到 {len(chapters)} 个章节：")

                for chap in chapters:
                    # 统计数据查询
                    try:
                        q_cnt = supabase.table("question_bank").select("id", count="exact").eq("chapter_id", chap[
                            'id']).execute().count
                    except:
                        q_cnt = 0
                    try:
                        m_cnt = supabase.table("materials").select("id", count="exact").eq("chapter_id",
                                                                                           chap['id']).execute().count
                    except:
                        m_cnt = 0

                    # 章节卡片
                    with st.expander(f"📑 {chap['title']} (题库: {q_cnt} | 教材片段: {m_cnt})"):
                        # 操作栏
                        c_op1, c_op2 = st.columns([1, 4])
                        with c_op1:
                            if st.button("🗑️ 清空数据", key=f"del_c_{chap['id']}",
                                         help="删除该章节下的所有题目和教材内容"):
                                supabase.table("materials").delete().eq("chapter_id", chap['id']).execute()
                                supabase.table("question_bank").delete().eq("chapter_id", chap['id']).execute()
                                st.toast("已清空该章节数据")
                                time.sleep(1)
                                st.rerun()
                        with c_op2:
                            st.caption(f"物理页码: P{chap['start_page']} - P{chap['end_page']}")

                        st.divider()

                        # === 🟢 优化后的预览逻辑 ===
                        preview_check = st.checkbox(f"👁️ 预览数据", key=f"view_mat_{chap['id']}")

                        if preview_check:
                            # 场景 A: 纯题库 (有题无教材)
                            if q_cnt > 0 and m_cnt == 0:
                                st.info(f"💡 检测到本章节包含 {q_cnt} 道题目，属于【习题库】。")
                                st.markdown("👉 请前往 **「🛠️ 数据管理 & 补录」** 板块进行题目的预览与可视编辑。")

                            # 场景 B: 教材 (有教材内容)
                            elif m_cnt > 0:
                                with st.spinner("正在拉取教材原文..."):
                                    try:
                                        mats = supabase.table("materials").select("content").eq("chapter_id",
                                                                                                chap['id']).order(
                                            "id").execute().data
                                        full_text = "\n\n".join([m['content'] for m in mats])

                                        st.caption(f"📊 教材原文概览 (共 {len(full_text)} 字)")
                                        st.text_area("原文快照 (只读)", value=full_text, height=300, disabled=True,
                                                     key=f"v_{chap['id']}")

                                        st.warning(
                                            "✏️ 如需修改原文(OCR纠错)，请前往 **「🛠️ 数据管理 & 补录」** > **「📘 教材内容修订」**。")
                                    except Exception as e:
                                        st.error(f"读取错误: {e}")

                            # 场景 C: 空章节
                            else:
                                st.warning("⚠️ 该章节暂无任何数据 (无题目也无教材)。")

# =========================================================
# 🎓 AI 课堂 (讲义) - V10.0 极简交互版 (一键补全+预览优先)
# =========================================================
elif menu == "🎓 AI 课堂 (讲义)":
    st.title("🎓 AI 深度课堂")
    st.caption("分步生成长篇讲义，支持断点续写、深度问答与实时编辑。")

    # --- 1. 选书选章 ---
    subjects = get_subjects()
    if not subjects: st.warning("请先去【资料库】初始化数据"); st.stop()

    c1, c2, c3 = st.columns(3)
    with c1:
        s_name = st.selectbox("科目", [s['name'] for s in subjects])
        sid = next(s['id'] for s in subjects if s['name'] == s_name)
    with c2:
        books = get_books(sid)
        bid = None
        if books:
            b_map = {b['title']: b['id'] for b in books}
            b_name = st.selectbox("书籍", list(b_map.keys()))
            bid = b_map[b_name]
    with c3:
        cid = None
        if bid:
            chaps = get_chapters(bid)
            if chaps:
                c_map = {c['title']: c['id'] for c in chaps}
                c_name = st.selectbox("章节", list(c_map.keys()), key="chap_selector")
                cid = c_map[c_name]

    if not cid:
        st.info("👈 请先在上方选择一个章节")
        st.stop()

    st.divider()

    # --- 2. 功能分区 ---
    tab_view, tab_gen = st.tabs(["📚 我的讲义本 (历史)", "✨ 分步生成工作台"])

    # ==========================================
    # Tab 1: 查看、修改、问答 (保持不变)
    # ==========================================
    with tab_view:
        try:
            lessons = supabase.table("ai_lessons").select("*").eq("chapter_id", cid).eq("user_id", user_id).order(
                "created_at", desc=True).execute().data
        except:
            lessons = []

        if not lessons:
            st.info("📭 本章节暂无讲义，请去“生成工作台”创建一个吧！")
        else:
            for les in lessons:
                les_id = les['id']
                with st.expander(f"📝 {les['title']}", expanded=False):
                    # 标题修改
                    c_edit_t, c_save_t = st.columns([4, 1])
                    with c_edit_t:
                        new_t = st.text_input("标题", value=les['title'], key=f"title_in_{les_id}",
                                              label_visibility="collapsed")
                    with c_save_t:
                        if new_t != les['title']:
                            if st.button("💾", key=f"save_t_{les_id}"):
                                supabase.table("ai_lessons").update({"title": new_t}).eq("id", les_id).execute()
                                st.rerun()

                    # 快捷工具栏
                    c_tts, c_del, c_export = st.columns([1, 1, 4])
                    with c_tts:
                        if st.button("🎧 朗读", key=f"tts_{les_id}"):
                            st.toast("正在生成语音...")
                            mp3 = generate_audio_file(les['content'][:3000])
                            if mp3: st.audio(mp3)
                    with c_del:
                        if st.button("🗑️ 删除", key=f"del_{les_id}"):
                            supabase.table("ai_lessons").delete().eq("id", les_id).execute()
                            st.rerun()

                    st.markdown("---")
                    st.markdown(les['content'])

                    # 简单问答
                    st.divider()
                    st.caption("💬 对本讲义提问")
                    if f"chat_{les_id}" not in st.session_state:
                        st.session_state[f"chat_{les_id}"] = les.get('chat_history') or []

                    for msg in st.session_state[f"chat_{les_id}"]:
                        with st.chat_message(msg['role']): st.write(msg['content'])

                    q_in = st.chat_input(f"基于本讲义提问...", key=f"in_{les_id}")
                    if q_in:
                        history = st.session_state[f"chat_{les_id}"]
                        history.append({"role": "user", "content": q_in})
                        prompt = f"【讲义内容】\n{les['content'][:10000]}\n\n【用户问题】{q_in}"
                        ans = call_ai_universal(prompt)
                        history.append({"role": "assistant", "content": ans})
                        supabase.table("ai_lessons").update({"chat_history": history}).eq("id", les_id).execute()
                        st.rerun()

    # ==========================================
    # Tab 2: 分步生成工作台 (V10.0: 极简交互版)
    # ==========================================
    with tab_gen:
        # 1. 准备教材数据
        mats = supabase.table("materials").select("content").eq("chapter_id", cid).execute().data
        if not mats:
            st.warning("⚠️ 本章节尚未上传教材资料，请先去【智能拆书】上传。")
        else:
            full_text = "\n".join([m['content'] for m in mats])
            total_len = len(full_text)

            # --- 核心状态管理 ---
            DRAFT_KEY = f"draft_content_{cid}_{user_id}"
            CURSOR_KEY = f"char_cursor_{cid}_{user_id}"
            OUTLINE_KEY = f"outline_{cid}_{user_id}"
            EDITOR_KEY = f"editor_widget_{cid}"
            OVERWRITE_KEY = f"overwrite_pending_{cid}"
            GEN_LOCK_KEY = f"gen_lock_{cid}"
            EDIT_MODE_KEY = f"is_editing_{cid}"  # 新增：控制是否处于编辑模式

            # 初始化
            if DRAFT_KEY not in st.session_state: st.session_state[DRAFT_KEY] = ""
            if CURSOR_KEY not in st.session_state: st.session_state[CURSOR_KEY] = 0
            if OUTLINE_KEY not in st.session_state: st.session_state[OUTLINE_KEY] = []
            if OVERWRITE_KEY not in st.session_state: st.session_state[OVERWRITE_KEY] = None
            if GEN_LOCK_KEY not in st.session_state: st.session_state[GEN_LOCK_KEY] = False
            if EDIT_MODE_KEY not in st.session_state: st.session_state[EDIT_MODE_KEY] = False

            # --- 2. 智能大纲 ---
            if not st.session_state[OUTLINE_KEY]:
                with st.expander("✨ 智能大纲 (点击生成)", expanded=True):
                    st.info("💡 系统将扫描教材生成核心考点地图。")
                    if st.button("🔍 分析本章考点"):
                        with st.spinner("AI 正在构建知识地图..."):
                            res = get_cached_outline_v2(cid, full_text, user_id)
                            st.session_state[OUTLINE_KEY] = res
                            st.rerun()

            # --- 3. 进度与可视化 ---
            curr_pos = st.session_state[CURSOR_KEY]
            outline_data = st.session_state[OUTLINE_KEY]
            current_draft = st.session_state[DRAFT_KEY]
            outline_status = check_outline_coverage_v2(outline_data, current_draft)

            # 顶部仪表盘
            c_p1, c_p2, c_p3 = st.columns(3)
            with c_p1:
                prog = min(curr_pos / total_len, 1.0)
                st.metric("📖 阅读进度", f"{int(prog * 100)}%")
                st.progress(prog)
            with c_p2:
                if outline_status:
                    covered = sum(1 for x in outline_status if x['covered'])
                    total_pts = len(outline_status)
                    st.metric("🗺️ 知识点覆盖", f"{covered}/{total_pts}")
                    st.progress(covered / total_pts if total_pts else 0)
                else:
                    st.metric("🗺️ 知识点", "--")
            with c_p3:
                CHUNK_SIZE = 3500
                step_len = CHUNK_SIZE - 200  # 实际每一步推进的距离
                remaining_chars = max(0, total_len - curr_pos)

                # 🟢 优化：使用向上取整，哪怕只剩 100 字也算 1 步
                if remaining_chars > 0:
                    rem_steps = math.ceil(remaining_chars / step_len)
                else:
                    rem_steps = 0

                st.metric("⏳ 预计剩余步数", f"约 {rem_steps} 步")

            # --- 4. 主控区域 (双栏布局) ---
            col_left, col_right = st.columns([1, 3])

            # >>> 左侧：大纲导航 + 一键补全 <<<
            with col_left:
                st.markdown("#### 📌 知识地图")

                # 筛选出未覆盖的
                missing_items = [item for item in outline_status if not item['covered']] if outline_status else []

                if outline_status:
                    for idx, item in enumerate(outline_status):
                        c_icon, c_txt = st.columns([1, 5])
                        is_covered = item['covered']
                        icon = "✅" if is_covered else "🔴"

                        with c_icon:
                            st.write(icon)
                        with c_txt:
                            if not is_covered:
                                st.markdown(f"**{item['title']}**")
                            else:
                                st.caption(item['title'])

                    st.markdown("---")

                    # === 🔥 新增：一键补全功能 ===
                    if missing_items:
                        st.caption(f"检测到 {len(missing_items)} 个未覆盖知识点")
                        if st.button("⚡ 一键补全所有红圈", type="primary",
                                     help="AI 将自动撰写所有缺失的知识点并追加到文末"):
                            st.session_state[GEN_LOCK_KEY] = True
                            progress_bar = st.progress(0)
                            status_txt = st.empty()

                            try:
                                # 1. 循环生成
                                for i, m_item in enumerate(missing_items):
                                    status_txt.text(f"正在补写：{m_item['title']} ({i + 1}/{len(missing_items)})...")

                                    patch_prompt = f"""
                                    【任务】针对知识点“{m_item['title']}”写一段补充讲义。
                                    【风格】幽默风趣，多用 Emoji (✨,💡)。
                                    【要求】直接输出正文，举一个简单的例子辅助理解。不要写“好的”等废话。
                                    """
                                    res = call_ai_universal(patch_prompt)
                                    if res:
                                        new_block = f"\n\n### ✨ 补充重点：{m_item['title']}\n{res}"
                                        st.session_state[DRAFT_KEY] += new_block
                                        st.session_state[EDITOR_KEY] = st.session_state[DRAFT_KEY]

                                    progress_bar.progress((i + 1) / len(missing_items))

                                # 2. 强制拉满进度
                                st.session_state[CURSOR_KEY] = total_len

                                # 3. === 🟢 核心新增：自动保存到数据库 ===
                                try:
                                    exist = supabase.table("ai_lessons").select("id").eq("title", lesson_title).eq(
                                        "chapter_id", cid).execute().data
                                    if exist:
                                        supabase.table("ai_lessons").update({
                                            "content": st.session_state[DRAFT_KEY],
                                            "ai_model": style,
                                            "updated_at": "now()"
                                        }).eq("id", exist[0]['id']).execute()
                                    else:
                                        supabase.table("ai_lessons").insert({
                                            "user_id": user_id, "chapter_id": cid,
                                            "title": lesson_title,
                                            "content": st.session_state[DRAFT_KEY],
                                            "ai_model": style
                                        }).execute()
                                    st.toast("⚡ 补全完成，已自动存档！")
                                except Exception as e:
                                    st.error(f"自动保存失败: {e}")
                                # ========================================

                                status_txt.success("✅ 所有红圈知识点已补全！")
                                time.sleep(1)
                                st.rerun()

                            except Exception as e:
                                st.error(f"补全过程中断: {e}")
                            finally:
                                st.session_state[GEN_LOCK_KEY] = False

                    st.markdown("---")
                    if st.button("🧹 重置进度"):
                        st.session_state[DRAFT_KEY] = ""
                        st.session_state[CURSOR_KEY] = 0
                        st.session_state[OVERWRITE_KEY] = None
                        st.rerun()
                else:
                    st.caption("暂无大纲")

            # >>> 右侧：预览优先 + 编辑切换 <<<
            with col_right:
                c_conf1, c_conf2 = st.columns([1, 2])
                with c_conf1:
                    style = st.selectbox("授课风格",
                                         ["👶 小白通俗版 (趣味Emoji)", "🦁 考霸冲刺版 (干货)", "⚖️ 法条深度版"],
                                         label_visibility="collapsed")
                with c_conf2:
                    lesson_title = st.text_input("讲义标题", value=f"深度解析：{c_name}", label_visibility="collapsed")

                st.markdown("### 📄 讲义预览")

                # --- 编辑模式切换逻辑 ---
                is_editing = st.session_state[EDIT_MODE_KEY]

                # 容器：头部工具栏
                c_tool_1, c_tool_2 = st.columns([5, 1])
                with c_tool_2:
                    if not is_editing:
                        if st.button("✏️ 编辑", key="btn_enter_edit", help="点击进入手动编辑模式"):
                            st.session_state[EDIT_MODE_KEY] = True
                            st.rerun()
                    else:
                        if st.button("✅ 完成", key="btn_exit_edit", type="primary"):
                            # 退出编辑模式时，内容已经在 on_change 里同步了，这里只需切换状态
                            st.session_state[EDIT_MODE_KEY] = False
                            st.rerun()

                # 容器：内容显示区
                content_container = st.container(border=True)
                with content_container:
                    # 分支 A: 编辑模式
                    if is_editing:
                        # 定义同步回调
                        def sync_editor_change():
                            st.session_state[DRAFT_KEY] = st.session_state[EDITOR_KEY]


                        st.text_area(
                            "编辑区域",
                            value=st.session_state[DRAFT_KEY],
                            height=600,
                            key=EDITOR_KEY,
                            on_change=sync_editor_change,
                            label_visibility="collapsed",
                            placeholder="AI 生成的内容将出现在这里..."
                        )
                        st.caption("💡 提示：修改内容会自动保存到草稿，点击右上角“完成”返回预览。")

                    # 分支 B: 预览模式 (默认)
                    else:
                        if st.session_state[DRAFT_KEY]:
                            st.markdown(st.session_state[DRAFT_KEY], unsafe_allow_html=True)
                        else:
                            st.info(
                                "👋 欢迎使用 AI 课堂！\n\n请点击下方的 **“🚀 开始生成”** 按钮，AI 将根据教材为您分段生成讲义。\n\n左侧若有 **红圈** (未覆盖知识点)，可点击 **“⚡ 一键补全”** 进行查漏补缺。")

                # --- 底部控制栏 (生成 & 保存) ---
                # 1. 定义备份用的 Key (用于撤销功能)
                BACKUP_DRAFT_KEY = f"backup_draft_{cid}_{user_id}"
                BACKUP_CURSOR_KEY = f"backup_cursor_{cid}_{user_id}"
                if BACKUP_DRAFT_KEY not in st.session_state:
                    st.session_state[BACKUP_DRAFT_KEY] = None
                    st.session_state[BACKUP_CURSOR_KEY] = 0

                start_idx = st.session_state[CURSOR_KEY]
                end_idx = min(start_idx + CHUNK_SIZE, total_len)

                is_all_covered = outline_status and all(item['covered'] for item in outline_status)
                is_finished = (start_idx >= total_len) or is_all_covered

                st.divider()

                # 定义布局：生成控制(含撤销) | 保存 | 下一章
                b_col1, b_col2, b_col3 = st.columns([2, 1, 1])

                # >>> 左侧：生成与撤销逻辑 <<<
                with b_col1:
                    if is_editing:
                        st.warning("⚠️ 请先点击右上角“完成”退出编辑模式。")
                    else:
                        if is_finished:
                            st.success("🎉 本章内容已生成完毕！")
                            if st.button("🎓 生成结语 (Auto-Save)", type="primary", use_container_width=True):
                                # ... (此处保持之前的结语生成逻辑，它已经包含了自动保存，无需修改) ...
                                # 为了节省篇幅，这里假设你已经用了上一步提供的结语自动保存代码
                                pass
                                # (请确保使用我上一次回复中提供的“结语自动保存”代码块)

                        else:
                            # 嵌套列：生成 | 撤销
                            gen_col, undo_col = st.columns([3, 2])

                            # >>> A. 生成按钮 (自动保存) <<<
                            with gen_col:
                                btn_txt = "🚀 开始生成" if start_idx == 0 else "➕ 继续生成下一节"
                                if not st.session_state[GEN_LOCK_KEY]:
                                    if st.button(btn_txt, type="primary", use_container_width=True):
                                        st.session_state[GEN_LOCK_KEY] = True
                                        try:
                                            # 备份
                                            st.session_state[BACKUP_DRAFT_KEY] = st.session_state[DRAFT_KEY]
                                            st.session_state[BACKUP_CURSOR_KEY] = st.session_state[CURSOR_KEY]

                                            emoji_instruct = "大量使用 Emoji (💡,✨,💰,⚠️) 使得排版活泼有趣。" if "小白" in style else "适当使用图标强调重点。"
                                            chunk_text = full_text[start_idx:end_idx]
                                            context_text = st.session_state[DRAFT_KEY][-800:] if len(
                                                st.session_state[DRAFT_KEY]) > 0 else ""

                                            prompt = f"""
                                            【角色】金牌会计讲师
                                            【风格】{style}
                                            【视觉要求】{emoji_instruct}
                                            【任务】讲解以下教材片段。
                                            【当前教材】{chunk_text}
                                            【上文回顾】...{context_text}
                                            【排版要求】
                                            1. 使用 Markdown 标题 (##, ###)。
                                            2. 重点概念加粗。
                                            3. **遇到难点必须举生活中的例子** (例如：买菜、谈恋爱、开公司)。
                                            """
                                            with st.spinner("AI 正在备课中..."):
                                                res = call_ai_universal(prompt)
                                                if res and "Error" not in res:
                                                    sep = "\n\n---\n\n" if start_idx > 0 else ""
                                                    updated_full = st.session_state[DRAFT_KEY] + sep + res

                                                    # 更新 Session
                                                    st.session_state[DRAFT_KEY] = updated_full
                                                    st.session_state[EDITOR_KEY] = updated_full
                                                    next_pos = max(end_idx - 200, start_idx + 100)
                                                    st.session_state[CURSOR_KEY] = min(next_pos, total_len)

                                                    # === 🟢 核心新增：生成后立即自动保存 ===
                                                    try:
                                                        exist = supabase.table("ai_lessons").select("id").eq("title",
                                                                                                             lesson_title).eq(
                                                            "chapter_id", cid).execute().data
                                                        if exist:
                                                            supabase.table("ai_lessons").update({
                                                                "content": updated_full, "ai_model": style,
                                                                "updated_at": "now()"
                                                            }).eq("id", exist[0]['id']).execute()
                                                        else:
                                                            supabase.table("ai_lessons").insert({
                                                                "user_id": user_id, "chapter_id": cid,
                                                                "title": lesson_title,
                                                                "content": updated_full, "ai_model": style
                                                            }).execute()
                                                        # 只有第一次生成才弹窗，避免每一步都弹窗打扰
                                                        if start_idx == 0: st.toast("💾 已自动建立存档")
                                                    except Exception as e:
                                                        print(f"Auto-save failed: {e}")
                                                    # ========================================

                                                else:
                                                    st.error(f"生成失败: {res}")
                                        finally:
                                            st.session_state[GEN_LOCK_KEY] = False
                                            st.rerun()
                                else:
                                    st.info("🔄 正在生成中...")

                            # >>> B. 撤销按钮 (撤销后也自动保存，保持数据库同步) <<<
                            with undo_col:
                                if st.session_state[BACKUP_DRAFT_KEY] is not None and st.session_state[DRAFT_KEY] != \
                                        st.session_state[BACKUP_DRAFT_KEY]:
                                    if st.button("↩️ 撤销本次", help="撤销刚才的操作（数据库也会回滚）",
                                                 use_container_width=True):
                                        # 还原状态
                                        st.session_state[DRAFT_KEY] = st.session_state[BACKUP_DRAFT_KEY]
                                        st.session_state[EDITOR_KEY] = st.session_state[BACKUP_DRAFT_KEY]
                                        st.session_state[CURSOR_KEY] = st.session_state[BACKUP_CURSOR_KEY]

                                        # === 🟢 核心新增：撤销后同步更新数据库 ===
                                        try:
                                            exist = supabase.table("ai_lessons").select("id").eq("title",
                                                                                                 lesson_title).eq(
                                                "chapter_id", cid).execute().data
                                            if exist:
                                                supabase.table("ai_lessons").update({
                                                    "content": st.session_state[DRAFT_KEY],  # 存入回滚后的内容
                                                    "updated_at": "now()"
                                                }).eq("id", exist[0]['id']).execute()
                                        except:
                                            pass
                                        # ========================================

                                        st.session_state[BACKUP_DRAFT_KEY] = None
                                        st.toast("已撤销并同步数据库 🔄")
                                        time.sleep(0.5)
                                        st.rerun()

                # 2. 中间：手动保存逻辑 (保留，供用户手动修改后保存)
                with b_col2:
                    final_content = st.session_state[DRAFT_KEY]
                    # 这里的逻辑不需要变，因为用户可能手动编辑了内容，需要一个显式的保存按钮
                    if st.session_state[OVERWRITE_KEY] is None:
                        if st.button("💾 手动保存", help="如果您手动编辑了内容，请点此保存", use_container_width=True):
                            if len(final_content) < 10:
                                st.warning("内容过少")
                            else:
                                # (此处省略重复的 Database Insert/Update 代码，保持原样即可)
                                # ... Database Save Logic ...
                                exist = supabase.table("ai_lessons").select("id").eq("title", lesson_title).eq(
                                    "chapter_id", cid).execute().data
                                if exist:
                                    supabase.table("ai_lessons").update({
                                        "content": final_content, "ai_model": style, "updated_at": "now()"
                                    }).eq("id", exist[0]['id']).execute()
                                else:
                                    supabase.table("ai_lessons").insert({
                                        "user_id": user_id, "chapter_id": cid, "title": lesson_title,
                                        "content": final_content, "ai_model": style
                                    }).execute()

                                st.balloons()
                                st.success("🎉 手动保存成功！")
                    else:
                        st.warning("⚠️ 文件已存在！")
                        if st.button("覆盖保存", type="primary"):
                            # ... 覆盖逻辑 ...
                            target_id = st.session_state[OVERWRITE_KEY]
                            supabase.table("ai_lessons").update({"content": final_content, "ai_model": style}).eq("id",
                                                                                                                  target_id).execute()
                            st.session_state[OVERWRITE_KEY] = None
                            st.toast("✅ 已覆盖")
                            time.sleep(1)
                            st.rerun()

                # 3. 右侧：👉 下一章
                with b_col3:
                    if is_finished:
                        all_chap_titles = list(c_map.keys())
                        try:
                            curr_idx = all_chap_titles.index(c_name)
                        except:
                            curr_idx = -1

                        if curr_idx != -1 and curr_idx < len(all_chap_titles) - 1:
                            next_chap_title = all_chap_titles[curr_idx + 1]
                            st.write("")
                            if st.button(f"➡️ 下一章", help=f"自动跳转至：{next_chap_title}", use_container_width=True):
                                st.session_state["chap_selector"] = next_chap_title
                                st.rerun()
                        else:
                            st.info("🏁 已是最后一章")
# =========================================================
# 📝 章节特训 (V6.3: 完整逻辑修复版 - 含数据库查询与主观题支持)
# =========================================================
elif menu == "📝 章节特训":
    st.title("📝 章节突破")

    # --- 1. JS 实时悬浮计时器 ---
    if st.session_state.get('quiz_active'):
        if 'js_start_time' not in st.session_state:
            st.session_state.js_start_time = int(time.time() * 1000)

        components.html(f"""
        <div style='position:fixed;top:60px;right:20px;z-index:9999;background:linear-gradient(45deg, #00C090, #00E6AC);color:white;padding:8px 20px;border-radius:30px;font-family:monospace;font-size:18px;font-weight:bold;box-shadow:0 4px 15px rgba(0,192,144,0.3)'>
            ⏱️ <span id='t'>00:00</span>
        </div>
        <script>
            setInterval(()=>{{
                var d=Math.floor((Date.now()-{st.session_state.js_start_time})/1000);
                var m=Math.floor(d/60).toString().padStart(2,'0');
                var s=(d%60).toString().padStart(2,'0');
                document.getElementById('t').innerText=m+':'+s;
            }},1000)
        </script>
        """, height=0)

    # --- 2. 启动配置区 (未开始状态) ---
    if not st.session_state.get('quiz_active'):
        subjects = get_subjects()
        if subjects:
            # 级联选择器
            c1, c2, c3 = st.columns(3)
            with c1:
                s_name = st.selectbox("1. 选择科目", [s['name'] for s in subjects])
                sid = next(s['id'] for s in subjects if s['name'] == s_name)

            with c2:
                books = get_books(sid)
                if not books:
                    st.warning("该科目下无书籍")
                    bid = None
                else:
                    b_map = {f"{b['title']} (ID:{b['id']})": b['id'] for b in books}
                    sel_b_label = st.selectbox("2. 选择书籍", list(b_map.keys()))
                    bid = b_map[sel_b_label]

            with c3:
                cid = None
                if bid:
                    chaps = get_chapters(bid)
                    if not chaps:
                        st.warning("本书无章节")
                    else:
                        c_map = {f"{c['title']} (ID:{c['id']})": c['id'] for c in chaps}
                        sel_c_label = st.selectbox("3. 选择章节", list(c_map.keys()))
                        cid = c_map[sel_c_label]

            # 选中章节后的逻辑
            if cid:
                st.markdown("---")

                # === 📊 智能进度看板 ===
                try:
                    q_res = supabase.table("question_bank").select("id").eq("chapter_id", cid).execute().data
                    total_q = len(q_res)
                    
                    mastered_count = 0
                    done_ids = []
                    if total_q > 0:
                        user_correct = supabase.table("user_answers").select("question_id").eq("user_id", user_id).eq("is_correct", True).execute().data
                        chapter_q_ids = set([q['id'] for q in q_res])
                        user_correct_ids = set([a['question_id'] for a in user_correct])
                        mastered_ids = user_correct_ids.intersection(chapter_q_ids)
                        mastered_count = len(mastered_ids)
                        done_ids = list(mastered_ids)
                    
                    prog = mastered_count / total_q if total_q > 0 else 0
                    st.caption(f"📈 掌握进度：{mastered_count} / {total_q} 题")
                    st.progress(prog)
                    
                except:
                    total_q = 0; done_ids = []

                st.divider()

                # === 🎯 练习模式选择 ===
                mode = st.radio("练习策略", [
                    "🧹 消灭库存 (只做未掌握的题)",
                    "🎲 随机巩固 (全库随机抽)",
                    "🧠 AI 基于教材出新题"
                ], horizontal=True)

                # 🔥 修改点 A：点击“开始练习”时清理旧数据
                if st.button("🚀 开始练习", type="primary", use_container_width=True):
                    # 1. 先彻底清理旧缓存
                    cleanup_quiz_session()

                    # --- 策略 A: 消灭库存 ---
                    if "消灭" in mode:
                        if total_q == 0:
                            st.error("题库为空，请先去【资料库】录入真题！")
                        elif mastered_count == total_q:
                            st.balloons()
                            st.success("🎉 本章题目已全部掌握！")
                        else:
                            # (此处使用之前修复过的稳健查询代码)
                            try:
                                if done_ids:
                                    ids_str = f"({','.join(map(str, done_ids))})"
                                    qs = supabase.table("question_bank").select("*").eq("chapter_id", cid).filter("id",
                                                                                                                  "not.in",
                                                                                                                  ids_str).limit(
                                        20).execute().data
                                else:
                                    qs = supabase.table("question_bank").select("*").eq("chapter_id", cid).limit(
                                        20).execute().data
                            except:
                                qs = []

                            if qs:
                                random.shuffle(qs)
                                st.session_state.quiz_data = qs[:10]
                                st.session_state.q_idx = 0
                                st.session_state.quiz_active = True
                                st.session_state.js_start_time = int(time.time() * 1000)
                                st.rerun()
                            else:
                                st.warning("数据加载异常，请重试")

                    # --- 策略 B: 随机巩固 ---
                    elif "随机" in mode:
                        if total_q == 0:
                            st.error("题库为空")
                        else:
                            qs = supabase.table("question_bank").select("*").eq("chapter_id", cid).limit(50).execute().data
                            if qs:
                                random.shuffle(qs)
                                st.session_state.quiz_data = qs[:10]
                                st.session_state.q_idx = 0
                                st.session_state.quiz_active = True
                                st.session_state.js_start_time = int(time.time() * 1000)
                                st.rerun()
                    
                    # --- 策略 C: AI 出题 ---
                    else:
                        mats = supabase.table("materials").select("content").eq("chapter_id", cid).execute().data
                        if not mats:
                            st.error("该章节没有上传教材资料！请去【资料库】上传 PDF/Word。")
                        else:
                            full_text = "\n".join([m['content'] for m in mats])
                            with st.spinner("🤖 AI 正在研读教材并出题..."):
                                prompt = f"""
                                请基于以下教材内容，生成 3 道选择题（含单选/多选）。
                                教材片段：{full_text[:10000]}
                                必须返回纯 JSON 列表格式：
                                [
                                  {{
                                    "content": "题目描述...",
                                    "options": ["A.选项1", "B.选项2", "C.选项3", "D.选项4"],
                                    "correct_answer": "AB", 
                                    "explanation": "详细解析...",
                                    "type": "multi"
                                  }}
                                ]
                                """
                                res = call_ai_universal(prompt)
                                if res:
                                    try:
                                        clean = res.replace("```json","").replace("```","").strip()
                                        s = clean.find('['); e = clean.rfind(']')+1
                                        d = json.loads(clean[s:e])
                                        
                                        # 存入数据库
                                        db_qs = [{
                                            'chapter_id': cid, 'user_id': user_id,
                                            'type': 'multi' if len(str(x.get('correct_answer','')))>1 else 'single',
                                            'content': x['content'],
                                            'options': x['options'],
                                            'correct_answer': x['correct_answer'],
                                            'explanation': x['explanation'],
                                            'origin': 'ai_gen',
                                            'batch_source': f'AI-{int(time.time())}'
                                        } for x in d]
                                        supabase.table("question_bank").insert(db_qs).execute()
                                        
                                        st.session_state.quiz_data = d
                                        st.session_state.q_idx = 0
                                        st.session_state.quiz_active = True
                                        st.session_state.js_start_time = int(time.time() * 1000)
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"AI 生成格式错误: {e}")
        else:
            st.warning("请先去【资料库】初始化科目和上传书籍")

    # --- 3. 做题交互界面 ---
    if st.session_state.get('quiz_active'):
        # === 🛡️ 安全检查：防止状态丢失导致的报错 ===
        if 'q_idx' not in st.session_state or 'quiz_data' not in st.session_state:
            st.session_state.quiz_active = False
            st.rerun()

        idx = st.session_state.q_idx
        data_len = len(st.session_state.quiz_data)
        
        if idx >= data_len:
            st.balloons()
            st.success("🎉 本轮练习完成！")
            if st.button("🔙 返回章节菜单"):
                st.session_state.quiz_active = False
                st.rerun()
        else:
            q = st.session_state.quiz_data[idx]
            
            # 顶部进度
            st.progress((idx + 1) / data_len)
            c_idx, c_end = st.columns([5, 1])
            with c_idx: st.caption(f"当前进度：{idx + 1} / {data_len}")
            with c_end:
                if st.button("🏁 结束"):
                    cleanup_quiz_session()  # <--- 调用清理
                    st.rerun()

            # 数据解析
            q_text = q.get('content') or q.get('question')
            q_type = q.get('type', 'single') # single, multi, subjective
            q_opts = q.get('options', [])
            std_ans = q.get('correct_answer') or q.get('answer')
            q_exp = q.get('explanation', '暂无解析')

            # 徽章显示
            badges = {
                "single": ("单选题", "#00C090"),
                "multi": ("多选题", "#ff9800"),
                "subjective": ("🧠 主观题", "#9c27b0")
            }
            b_label, b_color = badges.get(q_type, ("未知", "#888"))
            
            st.markdown(f"""
            <div class='css-card'>
                <div style="margin-bottom:10px">
                    <span style='background:{b_color};color:white;padding:2px 8px;border-radius:4px;font-size:12px'>{b_label}</span>
                </div>
                <h4 style="line-height:1.6; white-space: pre-wrap;">{q_text}</h4>
            </div>
            """, unsafe_allow_html=True)
            
            # --- 输入区域 ---
            user_val = ""
            sub_key = f"sub_state_{idx}"
            if sub_key not in st.session_state: st.session_state[sub_key] = False

            # 1. 主观题渲染
            if q_type == 'subjective':
                st.info("📝 请在下方输入你的计算过程或分录：")
                user_val = st.text_area("作答区", height=150, key=f"q_subj_{idx}", disabled=st.session_state[sub_key])
            
            # 2. 客观题渲染
            elif q_type == 'multi':
                st.caption("请选择所有正确选项（多选）：")
                selected = []
                for opt in q_opts:
                    if st.checkbox(opt, key=f"q_{idx}_{opt}", disabled=st.session_state[sub_key]):
                        selected.append(opt[0].upper())
                user_val = "".join(sorted(selected))
            else:
                st.caption("请选择唯一正确选项：")
                sel = st.radio("选项", q_opts, key=f"q_rad_{idx}", disabled=st.session_state[sub_key], label_visibility="collapsed")
                user_val = sel[0].upper() if sel else ""

            # --- 提交按钮 ---
            if st.button("✅ 提交答案", use_container_width=True) and not st.session_state[sub_key]:
                st.session_state[sub_key] = True
                st.rerun()

                # --- 判分与反馈 ---
                if st.session_state[sub_key]:
                    is_correct_bool = False
                    ai_feedback = ""

                    # A. 主观题：AI 评分
                    if q_type == 'subjective':
                        with st.spinner("🤖 AI 阅卷老师正在批改你的答案..."):
                            # 1. 评分并缓存 (保持存字典结构，不破坏存库逻辑)
                            grade_key = f"grade_res_{idx}"
                            if grade_key not in st.session_state:
                                grade_res = ai_grade_subjective(user_val, std_ans, q_text)
                                st.session_state[grade_key] = grade_res

                            # 2. 读取数据
                            res = st.session_state[grade_key]
                            score = res.get('score', 0)
                            ai_feedback = res.get('feedback', '')

                            # 3. 判定逻辑
                            is_correct_bool = (score >= 60)

                            # 4. UI 展示 (采纳 DeepSeek 的美化建议)
                            color = "#00C090" if score >= 80 else ("#ff9800" if score >= 60 else "#dc3545")
                            st.markdown(f"""
                                    <div style="padding:15px; background:{color}20; border-left:5px solid {color}; border-radius:5px; margin:10px 0;">
                                        <h3 style="color:{color}; margin:0">得分：{score} / 100</h3>
                                        <p style="margin-top:5px"><b>👩‍🏫 点评：</b>{ai_feedback}</p>
                                    </div>
                                    """, unsafe_allow_html=True)

                            # 5. 展示标准答案
                            with st.expander("查看参考答案"):
                                st.code(std_ans, language="markdown")

                # B. 客观题：逻辑匹配
                else:
                    clean_std = str(std_ans).replace(" ","").replace(",","").upper()
                    if user_val == clean_std:
                        st.markdown(f"<div class='success-box'>🎉 回答正确！</div>", unsafe_allow_html=True)
                        is_correct_bool = True
                    else:
                        st.error(f"❌ 遗憾答错。正确答案是：{clean_std}")
                        is_correct_bool = False
                    st.info(f"💡 **解析：** {q_exp}")

                # --- 存库逻辑 (通用) ---
                    # --- 存库逻辑 (V5.1 增强版：含分数与AI评价) ---
                    save_key = f"saved_db_{idx}"  # 利用 idx 生成唯一 Key，防止刷新页面重复插入

                    if save_key not in st.session_state:
                        try:
                            # 1. 检查题目 ID 是否存在
                            # (注：如果是 AI 临时生成的题目且未入库，q['id'] 可能为空，此时不存做题记录)
                            qid = q.get('id')

                            if qid:
                                # 2. 计算分数与反馈内容
                                final_score = 0
                                final_feedback = ""

                                if q_type == 'subjective':
                                    # === 主观题 ===
                                    # 从 Session 中获取之前 ai_grade_subjective 返回的结果
                                    grade_data = st.session_state.get(f"grade_res_{idx}", {})
                                    # 确保转为数字类型，防止 None
                                    final_score = float(grade_data.get('score', 0))
                                    final_feedback = str(grade_data.get('feedback', ''))
                                else:
                                    # === 客观题 (单选/多选) ===
                                    # 逻辑简单：对就是 100 分，错就是 0 分
                                    final_score = 100.0 if is_correct_bool else 0.0
                                    final_feedback = ""  # 客观题通常不需要 AI 评价，留空即可

                                # 3. 构造插入数据 Payload
                                payload = {
                                    "user_id": user_id,
                                    "question_id": qid,
                                    "user_response": user_val,  # 用户的原始作答
                                    "is_correct": is_correct_bool,  # 布尔值
                                    "score": final_score,  # 数值型分数
                                    "ai_feedback": final_feedback,  # AI 评语
                                    "exam_id": None  # 章节练习不属于模考，设为 Null
                                }

                                # 4. 执行数据库插入
                                supabase.table("user_answers").insert(payload).execute()

                                # 5. 标记为已保存 (关键步骤)
                                st.session_state[save_key] = True

                                # 可选：轻提示
                                # st.toast("💾 记录已保存")

                        except Exception as e:
                            # 捕获异常，防止因网络波动导致整个页面崩溃
                            print(f"❌ 存库失败 [QID: {q.get('id')}]: {e}")
            # 下一题
            st.divider()
            if st.button("➡️ 下一题", type="primary", use_container_width=True):
                if idx < data_len - 1:
                    st.session_state.q_idx += 1
                    st.rerun()
                else:
                    st.balloons()
                    st.success("🎉 本轮练习全部完成！")
                    if st.button("返回主菜单"):
                        cleanup_quiz_session()
                        st.rerun()


# =========================================================
# ⚔️ 全真模考 (V6.0: 混合题型 + 批量 AI 阅卷)
# =========================================================
elif menu == "⚔️ 全真模考":
    st.title("⚔️ 全真模拟考试")
    
    if 'exam_session' not in st.session_state:
        st.session_state.exam_session = None

    # 1. 配置台 (保持逻辑不变，只展示关键改动)
    if not st.session_state.exam_session:
        # ... (省略历史记录和选择科目代码) ...
        if st.button("🚀 生成试卷"):
             # ... (省略组卷逻辑，假设 final_paper 已生成) ...
             # 确保 final_paper 里包含了 subjective 类型的题
             st.session_state.exam_session = {
                 "paper": final_paper, # List of questions
                 "answers": {}, 
                 "submitted": False,
                 # ... 其他字段
             }
             st.rerun()

    # 2. 考试进行中
    elif not st.session_state.exam_session['submitted']:
        session = st.session_state.exam_session
        paper = session['paper']
        
        # ... (倒计时组件略) ...
        
        with st.form("exam_paper_form"):
            for idx, q in enumerate(paper):
                st.markdown(f"**第 {idx+1} 题**")
                
                # 渲染题干
                st.write(q['content'])
                
                q_type = q.get('type', 'single')
                
                # 区分渲染
                if q_type == 'subjective':
                    st.text_area("请输入答案/分录", key=f"ex_sub_{idx}", height=100)
                    # 注意：Form 里的 text_area 不需要 on_change，提交时会自动获取 session_state
                elif q_type == 'multi':
                    cols = st.columns(2)
                    for i, opt in enumerate(q.get('options',[])):
                        cols[i%2].checkbox(opt, key=f"ex_mul_{idx}_{i}")
                else:
                    st.radio("单选", q.get('options',[]), key=f"ex_sin_{idx}")
                
                st.divider()
            
            if st.form_submit_button("🏁 交卷"):
                # 收集答案
                for idx, q in enumerate(paper):
                    q_type = q.get('type', 'single')
                    if q_type == 'subjective':
                        session['answers'][idx] = st.session_state.get(f"ex_sub_{idx}", "")
                    elif q_type == 'multi':
                        sel = []
                        for i, opt in enumerate(q.get('options',[])):
                            if st.session_state.get(f"ex_mul_{idx}_{i}"): sel.append(opt[0].upper())
                        session['answers'][idx] = "".join(sorted(sel))
                    else:
                        val = st.session_state.get(f"ex_sin_{idx}")
                        session['answers'][idx] = val[0].upper() if val else ""
                
                session['submitted'] = True
                st.rerun()

    # 3. 考后报告 (含批量阅卷)
    else:
        session = st.session_state.exam_session
        paper = session['paper']
        user_ans_map = session['answers']
        
        # 如果还没出报告，先计算
        if 'report_data' not in session:
            total_score = 0
            detail_report = []
            
            # 创建进度条
            st.info("🤖 AI 正在逐题批改主观题，请稍候...")
            bar = st.progress(0)
            
            for idx, q in enumerate(paper):
                u_ans = user_ans_map.get(idx, "")
                q_type = q.get('type', 'single')
                std_ans = q.get('correct_answer', '')
                
                item_score = 0
                is_correct = False
                feedback = ""
                
                # 分支 A: 主观题 (调用 AI)
                if q_type == 'subjective':
                    res = ai_grade_subjective(u_ans, std_ans, q['content'])
                    # 假设每题权重平均，换算成百分制
                    # 比如试卷共10题，每题10分。AI给的 res['score'] 是0-100。
                    # 得分 = (res['score'] / 100) * (100 / len(paper))
                    weight = 100 / len(paper)
                    item_score = (res['score'] / 100) * weight
                    is_correct = (res['score'] >= 60)
                    feedback = res['feedback']
                
                # 分支 B: 客观题
                else:
                    weight = 100 / len(paper)
                    clean_std = str(std_ans).replace(" ","").upper()
                    if u_ans == clean_std:
                        item_score = weight
                        is_correct = True
                    else:
                        item_score = 0
                        is_correct = False
                
                total_score += item_score
                detail_report.append({
                    "q": q, "u_ans": u_ans, "score": item_score, 
                    "is_correct": is_correct, "feedback": feedback
                })
                bar.progress((idx+1)/len(paper))
            
            session['report_data'] = detail_report

            # === 🔥 新增：模考数据完整闭环入库 ===
            if 'saved_exam_flag' not in session:  # 防止刷新重复提交
                try:
                    # 1. 先创建模考记录 (Mock Exam Header)
                    exam_payload = {
                        "user_id": user_id,
                        "title": f"全真模考 {datetime.date.today()}",
                        "mode": "full",
                        "user_score": int(total_score),
                        "exam_data": detail_report  # 存快照，以防题目被删
                    }
                    exam_res = supabase.table("mock_exams").insert(exam_payload).execute()

                    # 获取新生成的 exam_id
                    new_exam_id = exam_res.data[0]['id']

                    # 2. 再批量插入做题详情，并关联 exam_id
                    db_answers = []
                    timestamp = datetime.datetime.now().isoformat()

                    for item in detail_report:
                        q_data = item['q']
                        # 确保只存已入库的题目的 ID
                        if q_data.get('id'):
                            db_answers.append({
                                "user_id": user_id,
                                "question_id": q_data.get('id'),
                                "exam_id": new_exam_id,  # <--- 关键：关联模考ID
                                "user_response": item['u_ans'],
                                "is_correct": item['is_correct'],
                                "score": item.get('score', 0),
                                "ai_feedback": item.get('feedback', ''),
                                "created_at": timestamp
                            })

                    if db_answers:
                        supabase.table("user_answers").insert(db_answers).execute()
                        st.toast(f"💾 模考存档成功！ID: {new_exam_id}")
                        session['saved_exam_flag'] = True

                except Exception as e:
                    st.error(f"模考存档失败: {e}")

            session['final_score'] = int(total_score)
            st.rerun()

        # 展示报告
        final_score = session['final_score']
        st.balloons()
        st.markdown(f"# 🏆 最终得分：{final_score}")
        
        for idx, item in enumerate(session['report_data']):
            q = item['q']
            status = "✅" if item['is_correct'] else "❌"
            with st.expander(f"第 {idx+1} 题 {status} (得分: {item['score']:.1f})"):
                st.write(q['content'])
                st.markdown(f"**你的答案：**\n{item['u_ans']}")
                st.markdown(f"**参考答案：**\n{q['correct_answer']}")
                if item['feedback']:
                    st.info(f"AI 点评：{item['feedback']}")
                    
        if st.button("退出"):
            st.session_state.exam_session = None
            st.rerun()

# === 📊 弱项分析 ===
elif menu == "📊 弱项分析":
    st.title("📊 学习效果")
    try:
        rows = supabase.table("user_answers").select("*").limit(500).execute().data
        if rows:
            df = pd.DataFrame(rows)
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"<div class='css-card'>总刷题<div class='stat-value'>{len(df)}</div></div>", unsafe_allow_html=True)
            with c2:
                ok = len(df[df['is_correct']==True])
                st.markdown(f"<div class='css-card'>正确率<div class='stat-value'>{int(ok/len(df)*100)}%</div></div>", unsafe_allow_html=True)
            
            fig = px.pie(df, names='is_correct', title='正确率分布', color_discrete_sequence=['#00C090', '#FF7043'])
            st.plotly_chart(fig)
        else: st.info("暂无数据")
    except: st.error("数据加载失败")


# =========================================================
# ❌ 错题本 (V8.1: 含删除/重生成功能的完整版)
# =========================================================
elif menu == "❌ 错题本":
    st.title("❌ 错题集 (智能私教版)")

    # 1. 获取错题数据
    try:
        errs = supabase.table("user_answers").select("*, question_bank(*)").eq("user_id", user_id).eq("is_correct",
                                                                                                      False).order(
            "created_at", desc=True).execute().data
    except Exception as e:
        st.error(f"数据加载失败: {e}")
        errs = []

    # 2. 去重逻辑
    uq = {}
    for e in errs:
        if not e['question_bank']: continue
        qid = e['question_id']
        if qid not in uq: uq[qid] = e

    if not uq:
        st.success("🎉 太棒了！目前没有待消灭的错题。")
    else:
        st.caption(f"共累计 {len(uq)} 道错题，加油消灭它们！")

        # 3. 遍历展示错题
        for qid, e in uq.items():
            q = e['question_bank']

            # --- 布局：题干区 ---
            with st.expander(f"🔴 [{q.get('type', '未知')}] {q['content'][:30]}...", expanded=False):
                # A. 题目详情
                st.markdown(f"### {q['content']}")
                q_type = q.get('type', 'single')

                # 选项字符串生成（用于 Prompt）
                options_text = ""
                if q_type in ['single', 'multi'] and q.get('options'):
                    st.markdown("---")
                    for o in q['options']:
                        st.markdown(f"<div class='option-item'>{o}</div>", unsafe_allow_html=True)
                    options_text = "\n".join([f"  {opt}" for opt in q['options']])
                else:
                    options_text = "（本题为主观题，无选项）"

                # B. 答案对比
                st.markdown("---")
                c1, c2 = st.columns(2)
                with c1:
                    st.error(f"❌ 你的答案：\n{e['user_response']}")
                with c2:
                    st.success(f"✅ 正确答案：\n{q['correct_answer']}")

                if q.get('explanation'):
                    with st.chat_message("assistant", avatar="📖"):
                        st.write(f"**参考解析：** {q['explanation']}")

                # --- AI 私教交互区 ---
                st.markdown("### 👩‍🏫 AI 私教辅导")

                # 读取历史对话
                chat_history = e.get('ai_chat_history') or []

                # === 🔥 修改点：带操作按钮的消息展示循环 ===
                # 我们需要修改 chat_history，所以用一个副本或标记来处理删除/重生成
                action_triggered = False  # 标记本次循环是否有操作，避免多次 rerun

                for idx, msg in enumerate(chat_history):
                    role = "user" if msg['role'] == "user" else "ai"
                    avatar = "🧑‍🎓" if role == "user" else "🤖"

                    with st.chat_message(role, avatar=avatar):
                        st.markdown(msg['content'])

                        # 仅为 AI 的回复添加操作按钮
                        if role == "ai":
                            c_del, c_regen, c_void = st.columns([1, 1, 6])

                            # 🗑️ 删除按钮
                            if c_del.button("🗑️", key=f"del_{qid}_{idx}", help="删除这条回答"):
                                chat_history.pop(idx)
                                # 如果删除的是回答，且上一条是用户的追问，通常也应该把上一条没用的追问删掉，
                                # 但为了灵活，这里只删当前这条。
                                supabase.table("user_answers").update({"ai_chat_history": chat_history}).eq("id", e[
                                    'id']).execute()
                                st.rerun()

                            # 🔄 重生成按钮
                            if c_regen.button("🔄", key=f"reg_{qid}_{idx}", help="对该回答不满意？重新生成"):
                                with st.spinner("🔄 AI 正在重写..."):
                                    # 1. 先删掉旧的
                                    chat_history.pop(idx)

                                    # 2. 确定 Prompt
                                    if idx == 0:
                                        # 情况 A: 这是第一条讲解。
                                        # 我们需要重建那个“超级 Prompt”。
                                        prompt = f"""
                                        【角色】幽默资深会计讲师。
                                        【任务】辅导错题。
                                        【题目】{q['content']}
                                        【选项】{options_text}
                                        【学生错解】{e['user_response']}
                                        【标准答案】{q['correct_answer']}
                                        【解析】{q.get('explanation', '无')}
                                        【要求】
                                        1. 🕵️ 诊断错因：为什么会选错？
                                        2. 💡 原理解析：大白话解释准则。
                                        3. 🍎 生活举例：必须举生活例子类比。
                                        """
                                        # 调用 AI (不带历史，因为这是第一条)
                                        new_reply = call_ai_universal(prompt, history=[])

                                    else:
                                        # 情况 B: 这是后续追问的回答。
                                        # 它的 Prompt 是上一条消息（idx-1）
                                        prev_user_msg = chat_history[idx - 1]['content']
                                        # 调用 AI (带上之前的历史作为上下文)
                                        # 注意：history 参数应该是 idx-1 之前的所有内容
                                        context_history = chat_history[:idx - 1]
                                        new_reply = call_ai_universal(prev_user_msg, history=context_history)

                                    # 3. 存入新回答
                                    if new_reply:
                                        chat_history.insert(idx, {"role": "model", "content": new_reply})
                                        supabase.table("user_answers").update({"ai_chat_history": chat_history}).eq(
                                            "id", e['id']).execute()
                                        st.rerun()

                # === 底部交互逻辑 (保持不变) ===
                c_act1, c_act2 = st.columns([1, 1])

                trigger_ai_first = False
                trigger_ai_followup = False
                user_question_text = ""

                # 场景 1: 还没聊过 (或第一条被删了)
                if not chat_history:
                    if c_act1.button("🙋‍♂️ 我没懂，请 AI 老师举例讲解", key=f"ai_teach_{qid}", type="primary"):
                        trigger_ai_first = True

                # 场景 2: 已经聊过
                else:
                    user_input = st.chat_input(f"追问 AI 老师 (ID: {qid})")
                    if user_input:
                        chat_history.append({"role": "user", "content": user_input})
                        supabase.table("user_answers").update({"ai_chat_history": chat_history}).eq("id",
                                                                                                    e['id']).execute()
                        trigger_ai_followup = True
                        user_question_text = user_input
                        st.rerun()  # 立即刷新显示用户的提问

                # === AI 生成逻辑执行区 ===
                if trigger_ai_first or (len(chat_history) > 0 and chat_history[-1]['role'] == 'user'):
                    # 这里加一个判断，只有当最后一条是 user 时才自动触发 AI，
                    # 避免页面刷新时意外触发
                    with st.spinner("🤖 AI 老师正在思考..."):
                        if not chat_history:  # 首次
                            prompt = f"""
                            【角色】幽默资深会计讲师。
                            【任务】辅导错题。
                            【题目】{q['content']}
                            【选项】{options_text}
                            【学生错解】{e['user_response']}
                            【标准答案】{q['correct_answer']}
                            【解析】{q.get('explanation', '无')}
                            【要求】
                            1. 🕵️ 诊断错因。
                            2. 💡 原理解析。
                            3. 🍎 生活举例（必选）。
                            """
                            reply = call_ai_universal(prompt, history=[])
                        else:  # 追问
                            last_q = chat_history[-1]['content']
                            # 传入除最后一条（也就是当前问题）之外的历史
                            reply = call_ai_universal(last_q, history=chat_history[:-1])

                        if reply:
                            chat_history.append({"role": "model", "content": reply})
                            supabase.table("user_answers").update({"ai_chat_history": chat_history}).eq("id", e[
                                'id']).execute()
                            st.rerun()

                # 移出按钮
                if c_act2.button("✅ 我学会了，移出", key=f"rm_{qid}"):
                    supabase.table("user_answers").update({"is_correct": True}).eq("id", e['id']).execute()
                    st.toast("已移出")
                    time.sleep(1);
                    st.rerun()



# =========================================================
# 🛠️ 数据管理 & 补录 (V7.0: 人工兜底与 Excel 导入)
# =========================================================
elif menu == "🛠️ 数据管理 & 补录":
    st.title("🛠️ 数据管理中心")
    st.caption("在此处手动修正 AI 的错误，或通过 Excel 批量导入自有题库。")

    # === 🟢 新增了中间的 Tab ===
    tab_edit_q, tab_edit_m, tab_upload = st.tabs(["✏️ 题库可视编辑", "📘 教材内容修订", "📥 Excel 批量导入"])


    # --- 公共选择器 (复用逻辑) ---
    def render_selectors(suffix):
        """渲染三级联动选择器，返回 (cid, c_name)"""
        subjects = get_subjects()
        if not subjects: st.warning("请先初始化科目"); return None, None

        c1, c2, c3 = st.columns(3)
        with c1:
            s_name = st.selectbox("科目", [s['name'] for s in subjects], key=f"sel_s_{suffix}")
            sid = next(s['id'] for s in subjects if s['name'] == s_name)
        with c2:
            books = get_books(sid)
            bid = None
            if books:
                b_map = {b['title']: b['id'] for b in books}
                b_name = st.selectbox("书籍", list(b_map.keys()), key=f"sel_b_{suffix}")
                bid = b_map[b_name]
        with c3:
            cid = None
            c_name = None
            if bid:
                chaps = get_chapters(bid)
                if chaps:
                    c_map = {c['title']: c['id'] for c in chaps}
                    c_name = st.selectbox("章节", list(c_map.keys()), key=f"sel_c_{suffix}")
                    cid = c_map[c_name]
        return cid, c_name


    # --- 辅助工具 ---
    def render_selectors(suffix, filter_mode=None):
        """
        渲染三级联动选择器
        :param filter_mode: None=不过滤; 'has_material'=只显示有教材原文的书
        :return: (cid, c_name)
        """
        subjects = get_subjects()
        if not subjects: st.warning("请先初始化科目"); return None, None

        c1, c2, c3 = st.columns(3)
        with c1:
            s_name = st.selectbox("科目", [s['name'] for s in subjects], key=f"sel_s_{suffix}")
            sid = next(s['id'] for s in subjects if s['name'] == s_name)
        with c2:
            # 1. 获取该科目下所有书
            all_books = get_books(sid)
            final_books = []

            # 2. 根据模式过滤书籍
            if not all_books:
                final_books = []
            elif filter_mode == "has_material":
                # === 🟢 核心过滤逻辑 ===
                # 只有当书籍下有关联的 materials 记录时才显示
                try:
                    # A. 拿到所有书的 ID
                    b_ids = [b['id'] for b in all_books]

                    # B. 查出这些书的所有章节
                    chaps_res = supabase.table("chapters").select("id, book_id").in_("book_id", b_ids).execute()
                    all_chaps = chaps_res.data

                    if all_chaps:
                        c_ids = [c['id'] for c in all_chaps]
                        # C. 查出这些章节中，哪些在 materials 表里有记录
                        # 这是一个轻量查询，只查 chapter_id
                        mat_res = supabase.table("materials").select("chapter_id").in_("chapter_id", c_ids).execute()
                        valid_chap_ids = set(m['chapter_id'] for m in mat_res.data)

                        # D. 反推有效的 book_id
                        valid_book_ids = set(c['book_id'] for c in all_chaps if c['id'] in valid_chap_ids)

                        # E. 过滤书籍列表
                        final_books = [b for b in all_books if b['id'] in valid_book_ids]
                    else:
                        final_books = []
                except Exception as e:
                    print(f"Filter Error: {e}")
                    final_books = all_books  # 降级处理：出错则显示全部
            else:
                final_books = all_books

            # 3. 渲染书籍下拉框
            bid = None
            if final_books:
                b_map = {b['title']: b['id'] for b in final_books}
                b_name = st.selectbox("书籍 (仅显示含教材)", list(b_map.keys()), key=f"sel_b_{suffix}")
                bid = b_map[b_name]
            else:
                st.warning("该科目下暂无符合要求的书籍")

        with c3:
            cid = None
            c_name = None
            if bid:
                chaps = get_chapters(bid)
                # 二次过滤：如果是编辑教材模式，只显示有教材的章节
                if filter_mode == "has_material":
                    # 同样的逻辑，确保选中的章节里真有东西
                    try:
                        c_ids_in_book = [c['id'] for c in chaps]
                        m_check = supabase.table("materials").select("chapter_id").in_("chapter_id",
                                                                                       c_ids_in_book).execute()
                        valid_c_ids = set(m['chapter_id'] for m in m_check.data)
                        chaps = [c for c in chaps if c['id'] in valid_c_ids]
                    except:
                        pass

                if chaps:
                    c_map = {c['title']: c['id'] for c in chaps}
                    c_name = st.selectbox("章节", list(c_map.keys()), key=f"sel_c_{suffix}")
                    cid = c_map[c_name]
                else:
                    st.caption("该书下无含教材的章节")
        return cid, c_name


    # --- 辅助工具 (保持不变) ---
    def list_to_str(lst):
        if isinstance(lst, list): return " | ".join(lst)
        return str(lst) if lst else ""


    def str_to_list(s):
        if not s: return []
        return [x.strip() for x in s.split("|") if x.strip()]


    # ---------------------------------------------------------
    # Tab 1: 题库可视编辑 (调用时 filter_mode=None)
    # ---------------------------------------------------------
    with tab_edit_q:
        st.info("💡 提示：双击单元格修改，修改后点击下方“💾 保存修改”生效。")
        cid_q, _ = render_selectors("q", filter_mode=None)  # 👈 题库不需要过滤

        if cid_q:
            # (...保持原有的题库编辑逻辑不变...)
            qs = supabase.table("question_bank").select("*").eq("chapter_id", cid_q).order("id").execute().data
            if not qs:
                st.warning("该章节暂无题目。")
            else:
                edit_data = []
                for q in qs:
                    edit_data.append({
                        "id": q['id'], "type": q['type'], "content": q['content'],
                        "options_str": list_to_str(q['options']), "correct_answer": q['correct_answer'],
                        "explanation": q.get('explanation', ''), "del": False
                    })
                df = pd.DataFrame(edit_data)
                edited_df = st.data_editor(
                    df,
                    column_config={
                        "id": st.column_config.NumberColumn("ID", disabled=True, width="small"),
                        "del": st.column_config.CheckboxColumn("删除?", width="small"),
                        "type": st.column_config.SelectboxColumn("题型", options=["single", "multi", "subjective"],
                                                                 width="medium"),
                        "content": st.column_config.TextColumn("题目内容", width="large"),
                        "options_str": st.column_config.TextColumn("选项", width="medium"),
                        "correct_answer": st.column_config.TextColumn("答案", width="small"),
                        "explanation": st.column_config.TextColumn("解析", width="medium"),
                    },
                    use_container_width=True, num_rows="dynamic", key=f"editor_q_{cid_q}"
                )
                if st.button("💾 保存题库修改", type="primary"):
                    try:
                        changes_count = 0
                        rows = edited_df.to_dict('records')
                        for row in rows:
                            if row.get('del') == True:
                                if row.get('id'): supabase.table("question_bank").delete().eq("id", row['id']).execute()
                                changes_count += 1
                                continue
                            clean_opts = str_to_list(row['options_str'])
                            payload = {
                                "chapter_id": cid_q, "user_id": user_id, "type": row['type'],
                                "content": row['content'], "options": clean_opts,
                                "correct_answer": row['correct_answer'], "explanation": row['explanation'],
                                "origin": "manual_edit"
                            }
                            if row.get('id'):
                                supabase.table("question_bank").update(payload).eq("id", row['id']).execute()
                            else:
                                if row['content']: supabase.table("question_bank").insert(payload).execute()
                            changes_count += 1
                        st.success(f"成功更新 {changes_count} 条记录！")
                        time.sleep(1);
                        st.rerun()
                    except Exception as e:
                        st.error(f"保存失败: {e}")

    # ---------------------------------------------------------
    # Tab 2: 教材内容修订 (调用时 filter_mode="has_material")
    # ---------------------------------------------------------
    with tab_edit_m:
        st.markdown("#### 📘 教材原文编辑器")
        st.caption("如果 AI 讲课时出现胡言乱语，通常是因为这里的**OCR原文识别错误**。请在此修正错别字。")

        # === 🟢 关键调用：启用 has_material 过滤模式 ===
        cid_m, c_name_m = render_selectors("m", filter_mode="has_material")

        if cid_m:
            try:
                mats = supabase.table("materials").select("*").eq("chapter_id", cid_m).order("id").execute().data
                if not mats:
                    # 理论上经过过滤不应该进这里，但为了保险
                    st.warning("⚠️ 该章节数据为空。")
                else:
                    mat_options = {f"片段 {i + 1} (ID: {m['id']}) - {m['content'][:20]}...": m for i, m in
                                   enumerate(mats)}
                    selected_label = st.selectbox("选择要编辑的片段", list(mat_options.keys()))
                    target_mat = mat_options[selected_label]

                    with st.form(key=f"edit_mat_form_{target_mat['id']}"):
                        new_content = st.text_area("编辑内容", value=target_mat['content'], height=400)
                        c_sub1, c_sub2 = st.columns([1, 5])
                        with c_sub1:
                            submit = st.form_submit_button("💾 保存修正", type="primary")
                        with c_sub2:
                            if st.form_submit_button("🗑️ 删除此片段"):
                                supabase.table("materials").delete().eq("id", target_mat['id']).execute()
                                st.toast("片段已删除");
                                time.sleep(1);
                                st.rerun()
                        if submit:
                            if new_content != target_mat['content']:
                                supabase.table("materials").update({"content": new_content}).eq("id", target_mat[
                                    'id']).execute()
                                st.success("✅ 教材内容已更新！");
                                time.sleep(1);
                                st.rerun()
                            else:
                                st.info("内容未变更。")
            except Exception as e:
                st.error(f"加载失败: {e}")

    # ---------------------------------------------------------
    # Tab 3: Excel 批量导入 (升级版：支持 Excel题库 + Word/Txt教材)
    # ---------------------------------------------------------
    with tab_upload:
        st.markdown("#### 📥 通用数据导入中心")
        st.caption("支持导入整理好的 **题库(Excel)** 或 **纯文本教材(Word/Txt)**。")

        # 1. 第一层：选择数据类型
        import_type = st.radio(
            "1. 请选择要导入的数据类型：",
            ["📝 题库数据 (Excel/CSV)", "📘 教材原文 (Word/Txt)"],
            horizontal=True
        )

        st.divider()

        # 2. 第二层：选择目标位置 (现有 vs 新建)
        target_mode = st.radio(
            "2. 请选择导入目标：",
            ["📂 导入到现有章节 (追加)", "🆕 新建书籍/章节 (新开)"],
            horizontal=True
        )

        final_cid = None  # 最终的目标章节ID
        final_c_name = ""  # 用于显示的章节名

        # === 逻辑分支 A: 现有章节 ===
        if "现有" in target_mode:
            # 这里调用选择器时，filter_mode=None，确保显示所有书（无论有没有教材）
            # 因为你要导入数据，所以应该允许往任何书里导
            cid_exist, c_name_exist = render_selectors("upload_exist", filter_mode=None)

            if cid_exist:
                final_cid = cid_exist
                final_c_name = c_name_exist
            else:
                st.info("👈 请在上方选择目标章节。")

        # === 逻辑分支 B: 新建书籍/章节 ===
        else:
            c_new1, c_new2, c_new3 = st.columns(3)
            with c_new1:
                # 必须选科目
                subjects = get_subjects()
                s_names = [s['name'] for s in subjects] if subjects else []
                sel_s_new = st.selectbox("所属科目", s_names, key="new_import_sub")
                # 找到 sid
                sel_sid = next((s['id'] for s in subjects if s['name'] == sel_s_new), None) if subjects else None

            with c_new2:
                new_book_title = st.text_input("新书籍名称", placeholder="例如：2026中级经济法-考前押题")

            with c_new3:
                new_chap_title = st.text_input("新章节名称", value="全文", placeholder="例如：第一章 总论")

            # 预校验
            if sel_sid and new_book_title and new_chap_title:
                st.success(f"准备创建：📘 {new_book_title} > 📑 {new_chap_title}")
                # 此时我们还没有 cid，需要在点击按钮那一刻创建
            else:
                st.warning("请补全书籍和章节名称。")

        st.divider()

        # 3. 第三层：上传与执行区

        # >>>>>>>> 场景 1: 题库导入 <<<<<<<<
        if "题库" in import_type:
            c_d1, c_d2 = st.columns([1, 2])
            with c_d1:
                st.markdown("**模板下载**")
                template_data = [{"题型(必填)": "single", "题目内容(必填)": "...", "选项(用|分隔)": "A.x | B.y",
                                  "正确答案(必填)": "A", "解析": ""}]
                df_temp = pd.DataFrame(template_data)
                csv = df_temp.to_csv(index=False).encode('utf-8-sig')
                st.download_button("⬇️ 题库模板.csv", data=csv, file_name="题库模板.csv", mime="text/csv")

            with c_d2:
                up_excel = st.file_uploader("上传 CSV/Excel", type=["csv", "xlsx"], key="up_q_bank")

            # 只有当 (选择了现有章节 OR 填写了新建信息) AND 上传了文件 时，按钮才可用
            ready_to_import = up_excel is not None
            if "现有" in target_mode and not final_cid: ready_to_import = False
            if "新建" in target_mode and (not new_book_title or not new_chap_title): ready_to_import = False

            if ready_to_import:
                if st.button("🚀 开始导入题库", type="primary"):
                    try:
                        # --- 如果是新建模式，先创建 DB 记录 ---
                        if "新建" in target_mode:
                            # 1. 建书
                            b_res = supabase.table("books").insert({
                                "user_id": user_id, "subject_id": sel_sid,
                                "title": new_book_title, "total_pages": 0
                            }).execute()
                            new_bid = b_res.data[0]['id']
                            # 2. 建章
                            c_res = supabase.table("chapters").insert({
                                "user_id": user_id, "book_id": new_bid,
                                "title": new_chap_title, "start_page": 0, "end_page": 0
                            }).execute()
                            final_cid = c_res.data[0]['id']
                            final_c_name = new_chap_title

                        # --- 开始读取文件 ---
                        if up_excel.name.endswith('.csv'):
                            df_new = pd.read_csv(up_excel)
                        else:
                            df_new = pd.read_excel(up_excel)

                        bar = st.progress(0)
                        batch_data = []
                        for i, row in df_new.iterrows():
                            # (数据清洗逻辑保持不变)
                            content = row.get('题目内容(必填)') or row.get('题目内容') or row.get('content')
                            ans = row.get('正确答案(必填)') or row.get('正确答案') or row.get('correct_answer')
                            if pd.isna(content) or pd.isna(ans): continue

                            opts_raw = row.get('选项(用|分隔)') or row.get('选项') or row.get('options')
                            opts_list = [str(x).strip() for x in str(opts_raw).split("|") if
                                         str(x).strip()] if opts_raw else []

                            batch_data.append({
                                "chapter_id": final_cid, "user_id": user_id,
                                "type": str(row.get('题型(必填)', 'single')).strip(),
                                "content": str(content), "correct_answer": str(ans),
                                "explanation": str(row.get('解析', '')), "options": opts_list,
                                "origin": "excel_import", "batch_source": f"Upload-{datetime.date.today()}"
                            })
                            if len(batch_data) >= 10:
                                supabase.table("question_bank").insert(batch_data).execute()
                                batch_data = []
                            bar.progress((i + 1) / len(df_new))

                        if batch_data: supabase.table("question_bank").insert(batch_data).execute()

                        st.balloons()
                        st.success(f"🎉 成功导入 {len(df_new)} 道题目至：{final_c_name}")
                        time.sleep(2);
                        st.rerun()

                    except Exception as e:
                        st.error(f"导入失败: {e}")

        # >>>>>>>> 场景 2: 教材导入 <<<<<<<<
        else:
            st.info("💡 提示：Word/Txt 内容将作为**教材片段**存入数据库。")
            up_doc = st.file_uploader("上传 Word (.docx) 或 文本 (.txt)", type=["docx", "txt"], key="up_doc_mat")

            ready_to_import = up_doc is not None
            if "现有" in target_mode and not final_cid: ready_to_import = False
            if "新建" in target_mode and (not new_book_title or not new_chap_title): ready_to_import = False

            if ready_to_import:
                if st.button("🚀 开始导入教材", type="primary"):
                    try:
                        # --- 如果是新建模式，先创建 DB 记录 ---
                        if "新建" in target_mode:
                            b_res = supabase.table("books").insert({
                                "user_id": user_id, "subject_id": sel_sid,
                                "title": new_book_title, "total_pages": 0
                            }).execute()
                            new_bid = b_res.data[0]['id']
                            c_res = supabase.table("chapters").insert({
                                "user_id": user_id, "book_id": new_bid,
                                "title": new_chap_title, "start_page": 0, "end_page": 0
                            }).execute()
                            final_cid = c_res.data[0]['id']
                            final_c_name = new_chap_title

                        # --- 解析文件 ---
                        content_extracted = ""
                        with st.spinner("正在解析文件..."):
                            if up_doc.name.endswith('.txt'):
                                content_extracted = up_doc.read().decode("utf-8")
                            elif up_doc.name.endswith('.docx'):
                                doc = docx.Document(up_doc)
                                content_extracted = "\n".join([p.text for p in doc.paragraphs])

                            if len(content_extracted) < 10:
                                st.error("❌ 文件内容过少，无法导入。")
                            else:
                                clean_content = clean_textbook_content(content_extracted)
                                save_material_v3(final_cid, clean_content, user_id)

                                st.balloons()
                                st.success(f"🎉 教材导入成功！已存入：{final_c_name}")
                                time.sleep(2);
                                st.rerun()

                    except Exception as e:
                        st.error(f"导入失败: {e}")



# =========================================================
# ⚙️ 设置中心 (V3.1 修复版：配置回显 + 连通测试 + 考期同步)
# =========================================================
elif menu == "⚙️ 设置中心":
    st.title("⚙️ 系统设置")
    
    # 1. 读取云端配置 (核心修复)
    # 必须先从 profile 里读出来，否则滑块永远是默认值
    current_settings = profile.get('settings') or {}
    saved_timeout = current_settings.get('ai_timeout', 60) # 读不到就默认60
    
    # --- A. AI 模型参数 ---
    st.markdown("#### 🤖 AI 参数配置")
    
    col_test, col_set = st.columns([1, 2])
    
    with col_test:
        st.info(f"当前通道：**{st.session_state.get('selected_provider')}**")
        # 连通性测试功能
        if st.button("📡 测试 AI 连通性"):
            with st.spinner("发送 Hello World..."):
                start_t = time.time()
                res = call_ai_universal("Say 'OK' in one word.", timeout_override=10)
                cost_t = time.time() - start_t
                
                if "Error" in res or "异常" in res:
                    st.error(f"❌ 失败: {res}")
                else:
                    st.success(f"✅ 通畅! 耗时 {cost_t:.2f}s")
                    st.caption(f"回复: {res}")

    with col_set:
        # 修复：value 设置为 saved_timeout (从数据库读)
        new_timeout = st.slider(
            "⏳ AI 回答超时限制 (秒)", 
            min_value=10, 
            max_value=300, 
            value=saved_timeout, 
            help="如果是生成整章讲义或全量入库，建议调大此值 (如 120秒)"
        )
        
        if st.button("💾 保存参数"):
            if new_timeout != saved_timeout:
                update_settings(user_id, {"ai_timeout": new_timeout})
                st.success(f"已保存！超时限制改为 {new_timeout} 秒")
                time.sleep(1)
                st.rerun() # 强制刷新页面，让变量生效
            else:
                st.info("配置未变更")

    st.divider()
    
    # --- B. 考试时间设置 (保留联网功能) ---
    st.markdown("#### 📅 考试倒计时")
    
    # 自动同步
    if st.button("🌐 联网推测 2025 考试日期 (AI)"):
        with st.spinner("正在分析历史考情..."):
            # 模拟 AI 决策
            p = f"根据中国会计资格评价中心惯例，推测 {datetime.date.today().year} 年中级会计考试日期。仅返回 YYYY-MM-DD 格式。"
            ai_date = call_ai_universal(p)
            try:
                clean_d = ai_date.strip()[:10]
                # 简单校验格式
                datetime.datetime.strptime(clean_d, '%Y-%m-%d')
                
                supabase.table("study_profile").update({"exam_date": clean_d}).eq("user_id", user_id).execute()
                st.success(f"已更新为: {clean_d}")
                time.sleep(1)
                st.rerun()
            except:
                st.warning("AI 返回日期格式有误，请手动设置")

    # 手动设置
    curr_date = datetime.date(2025, 9, 6)
    if profile.get('exam_date'):
        try: curr_date = datetime.datetime.strptime(profile['exam_date'], '%Y-%m-%d').date()
        except: pass
        
    set_date = st.date_input("手动设定目标日期", curr_date)
    if set_date != curr_date:
        supabase.table("study_profile").update({"exam_date": str(set_date)}).eq("user_id", user_id).execute()
        st.toast("日期已更新")
        time.sleep(1)
        st.rerun()

    st.divider()
    
    # --- C. 危险区域 ---
    with st.expander("🗑️ 危险操作 (数据清理)"):
        c_del1, c_del2 = st.columns(2)
        with c_del1:
            if st.button("清空所有错题记录"):
                supabase.table("user_answers").delete().eq("user_id", user_id).execute()
                st.success("错题本已清空")
                
        with c_del2:
            if st.button("清空所有书籍资料"):
                supabase.table("books").delete().eq("user_id", user_id).execute()
                # 因为设置了级联删除(Cascade)，章节、题目、内容会自动删除
                st.success("资料库已格式化")

